# rfp_processor.py
import os
import re
import openai
from django.conf import settings
from pathlib import Path
from docx import Document
import pdfplumber
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as LangchainDocument
# Build vector store from PDF files in knowledge directory
import hashlib
import pickle
import time
from datetime import datetime
import json
import pytesseract
from PIL import Image
import numpy as np
import cv2
import io

# Ensure OpenAI API key is set
openai.api_key = settings.OPENAI_API_KEY
os.environ["OPENAI_API_KEY"] = settings.OPENAI_API_KEY


# 🔹 Function to fix Arabic text
def fix_arabic_text(text):
    """Fix the direction of Arabic text extracted from PDF."""
    return text[::-1]


# 🆕 Function to detect if OCR is needed for a PDF
def is_ocr_needed(pdf_path):
    """
    Check if a PDF requires OCR by attempting to extract text directly first.
    Returns True if OCR is needed, False otherwise.
    """
    try:
        # Try to extract text directly
        with pdfplumber.open(pdf_path) as pdf:
            text_content = ""
            # Check first few pages only for efficiency
            max_pages_to_check = min(5, len(pdf.pages))
            for i in range(max_pages_to_check):
                page = pdf.pages[i]
                text = page.extract_text()
                if text and len(text.strip()) > 100:  # If we get reasonable text content
                    text_content += text

            # If we got less than expected text, OCR might be needed
            if len(text_content) < 50 * max_pages_to_check:
                return True
            return False
    except Exception as e:
        print(f"Error checking PDF for OCR: {str(e)}")
        return True  # Default to OCR if extraction fails


# 🆕 Function to perform OCR on a PDF
def perform_ocr_on_pdf(pdf_path, language='ara+eng'):
    """
    Perform OCR on a PDF file and return the extracted text.
    Optimized for Arabic and English text by default.
    """
    extracted_text = ""
    try:
        # Open PDF with PyMuPDF for better image extraction
        doc = fitz.open(pdf_path)
        total_pages = doc.page_count

        for page_num in range(total_pages):
            page = doc[page_num]

            # Get page as an image at higher resolution for better OCR
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_bytes = pix.tobytes("png")

            # Load image with PIL
            img = Image.open(io.BytesIO(img_bytes))

            # Convert to OpenCV format for preprocessing
            img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)

            # Preprocessing for better OCR results
            # Convert to grayscale
            gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)

            # Apply thresholding to handle different lighting conditions
            _, binary = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)

            # Noise removal
            kernel = np.ones((1, 1), np.uint8)
            opening = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel)

            # Convert back to PIL Image for Tesseract
            img_processed = Image.fromarray(cv2.bitwise_not(opening))

            # Perform OCR with tesseract, optimized for Arabic
            text = pytesseract.image_to_string(
                img_processed,
                lang=language,
                config='--psm 6 --oem 3'
            )

            if text.strip():
                # Fix Arabic text direction if needed
                text = fix_arabic_text(text)
                extracted_text += text + "\n\n"

            print(f"Processed page {page_num + 1}/{total_pages} with OCR")

        return extracted_text
    except Exception as e:
        print(f"Error performing OCR on PDF: {str(e)}")
        return ""


def extract_text_from_pdf(pdf_path, language='ara+eng', force_ocr=False):
    """
    Extract text from PDF with OCR as fallback if needed.
    Added force_ocr parameter to bypass detection when needed.
    """
    try:
        # Skip detection if OCR is forced
        if force_ocr:
            print(f"🔍 Forcing OCR processing: {pdf_path}")
            return perform_ocr_on_pdf(pdf_path, language)

        # First check if OCR is needed
        if is_ocr_needed(pdf_path):
            print(f"🔍 PDF requires OCR processing: {pdf_path}")
            return perform_ocr_on_pdf(pdf_path, language)

        # If OCR not needed, use regular extraction
        print(f"📄 Extracting text directly from PDF: {pdf_path}")
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += fix_arabic_text(extracted_text) + "\n"

        # Double-check if we got enough text
        if len(text.strip()) < 200:
            print(f"⚠️ Direct extraction yielded insufficient text. Falling back to OCR.")
            return perform_ocr_on_pdf(pdf_path, language)

        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        # Fallback to OCR if regular extraction fails
        return perform_ocr_on_pdf(pdf_path, language)


def save_rfp_sections_to_word(sections, output_path, competition_name):
    """
    Save the RFP sections to a Word document with proper RTL formatting and improved table support.
    Enhanced version with better paragraph handling and formatting.
    """
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt, RGBColor, Cm
    import re
    import os

    document = Document()

    # Set document margins
    for section in document.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # Add document title
    title = document.add_heading(competition_name, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x22, 0x35, 0x1D)

    # Set default paragraph style
    style = document.styles['Normal']
    style.font.rtl = True
    style.font.size = Pt(13)
    style.font.name = 'Traditional Arabic'
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    style.paragraph_format.space_after = Pt(8)  # Add space after paragraphs

    # Compile regular expressions for pattern matching
    table_start_pattern = re.compile(r'^\s*\|.*\|\s*$')
    table_header_separator = re.compile(r'^\s*\|?\s*[-:]+\s*\|[-:\s]*\|?\s*$')
    heading_pattern = re.compile(r'^#+\s+(.+)$')
    bullet_pattern = re.compile(r'^[\u2022\-*]\s+(.+)$')
    numbered_pattern = re.compile(r'^(\d+)\.\s+(.+)$')
    section_title_patterns = {}

    # Arabic ordinal numbers for sections
    section_number_words = ["الأول", "الثاني", "الثالث", "الرابع", "الخامس",
                            "السادس", "السابع", "الثامن", "التاسع", "العاشر",
                            "الحادي عشر", "الثاني عشر"]

    # Preprocess sections to remove unwanted formatting
    for i, (section_title, section_content) in enumerate(sections):
        sections[i] = (section_title, section_content.replace('**', ''))
        patterns = [
            f"القسم {section_number_words[i]}: {section_title}",
            f"القسم {i + 1}: {section_title}",
            f"{section_title}",
            f"{i + 1}. {section_title}",
            f"{i + 1}- {section_title}"
        ]
        section_title_patterns[i] = [re.compile(pattern) for pattern in patterns]

    # Process each section
    for section_index, (section_title, section_content) in enumerate(sections, 1):
        # Add section heading
        formatted_title = f"القسم {section_number_words[section_index - 1]}: {section_title}"
        heading = document.add_heading(level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = heading.add_run(formatted_title)
        run.font.rtl = True
        run.font.size = Pt(15)
        run.font.name = 'Traditional Arabic'
        run.font.color.rgb = RGBColor(0x22, 0x35, 0x1D)

        # Split section content into lines
        section_content_lines = section_content.split('\n')

        # Find where to start processing (skip section title if present in content)
        start_line = 0
        for i, line in enumerate(section_content_lines):
            if any(pattern.search(line) for pattern in section_title_patterns[section_index - 1]):
                start_line = i + 1
            else:
                break

        lines = section_content_lines[start_line:]
        current_paragraph = None
        table_rows = []
        collecting_table = False
        i = 0

        # Process each line
        while i < len(lines):
            line = lines[i].rstrip()

            # Handle table lines
            if table_start_pattern.match(line):
                if not collecting_table:
                    collecting_table = True
                    table_rows = []

                cleaned = line.strip('| \t')
                table_rows.append([cell.strip() for cell in cleaned.split('|')])

                if i + 1 < len(lines) and table_header_separator.match(lines[i + 1]):
                    i += 1

            # Finalize and add table
            elif collecting_table:
                collecting_table = False
                if table_rows:
                    num_cols = max(len(row) for row in table_rows)
                    if num_cols > 0 and len(table_rows) > 0:
                        table = document.add_table(rows=len(table_rows), cols=num_cols)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.RIGHT

                        for row_idx, row_cells in enumerate(table_rows):
                            row = table.rows[row_idx]
                            for col_idx, cell_text in enumerate(row_cells):
                                if col_idx < num_cols:
                                    cell = table.cell(row_idx, col_idx)
                                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                    p = cell.paragraphs[0]
                                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    try:
                                        p._element.get_or_add_pPr().set('bidi', '1')
                                    except:
                                        pass

                                    run = p.add_run(cell_text)
                                    run.font.rtl = True
                                    run.font.name = 'Traditional Arabic'
                                    run.font.size = Pt(13)

                        # Add extra space after table
                        document.add_paragraph()

            # Handle headings
            elif heading_match := heading_pattern.match(line):
                level = line.count('#')
                text = heading_match.group(1).strip()
                if level <= 4:
                    h = document.add_heading(level=level + 1)
                    h.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    run = h.add_run(text)
                    run.font.rtl = True
                    run.font.size = Pt(15)
                    run.font.name = 'Traditional Arabic'
                    run.font.color.rgb = RGBColor(0x22, 0x35, 0x1D)
                else:
                    p = document.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    p.paragraph_format.space_after = Pt(8)  # Add space after paragraph
                    run = p.add_run(text)
                    run.bold = True
                    run.font.rtl = True
                    run.font.size = Pt(13)
                    run.font.name = 'Traditional Arabic'
                current_paragraph = None  # Reset current paragraph

            # Handle bullet points
            elif bullet_match := bullet_pattern.match(line):
                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                p.paragraph_format.left_indent = Pt(12)
                p.paragraph_format.space_after = Pt(8)  # Add space after paragraph
                try:
                    p._element.get_or_add_pPr().set('bidi', '1')
                except:
                    pass
                bullet_text = bullet_match.group(1).strip()
                run = p.add_run('• ')
                run.font.rtl = True
                run.font.name = 'Traditional Arabic'
                run.font.size = Pt(13)
                run = p.add_run(bullet_text)
                run.font.rtl = True
                run.font.name = 'Traditional Arabic'
                run.font.size = Pt(13)
                current_paragraph = None  # Reset current paragraph

            # Handle numbered list items
            elif numbered_match := numbered_pattern.match(line):
                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                p.paragraph_format.left_indent = Pt(12)
                p.paragraph_format.space_after = Pt(8)  # Add space after paragraph
                try:
                    p._element.get_or_add_pPr().set('bidi', '1')
                except:
                    pass
                number = numbered_match.group(1)
                text = numbered_match.group(2).strip()
                run = p.add_run(f"{number}. ")
                run.font.rtl = True
                run.font.name = 'Traditional Arabic'
                run.font.size = Pt(13)
                run = p.add_run(text)
                run.font.rtl = True
                run.font.name = 'Traditional Arabic'
                run.font.size = Pt(13)
                current_paragraph = None  # Reset current paragraph

            # Handle regular text
            elif line.strip():
                # Start a new paragraph only if needed
                if not current_paragraph:
                    current_paragraph = document.add_paragraph()
                    current_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    current_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    current_paragraph.paragraph_format.space_after = Pt(8)  # Add space after paragraph

                # Check if we need to add a line break within the paragraph
                if current_paragraph.runs and i > 0 and lines[i - 1].strip():
                    current_paragraph.add_run(' ')  # Add space between sentences

                run = current_paragraph.add_run(line)
                run.font.rtl = True
                run.font.size = Pt(13)
                run.font.name = 'Traditional Arabic'
            else:
                # Empty line - end current paragraph
                current_paragraph = None

            i += 1

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    document.save(output_path)
    return output_path


def build_vector_store(knowledge_dir):
    """
    Process PDF files in the knowledge directory and build a vector store.
    Uses caching to avoid rebuilding if files haven't changed.

    Args:
        knowledge_dir: Directory containing the RFP PDF files

    Returns:
        FAISS vector store object
    """
    cache_dir = os.path.join(knowledge_dir, ".cache")
    os.makedirs(cache_dir, exist_ok=True)

    # Create a hash of the PDF files to check if they've changed
    pdf_files = [os.path.join(knowledge_dir, f) for f in os.listdir(knowledge_dir) if f.endswith('.pdf')]
    pdf_files.sort()  # Ensure consistent ordering

    # Get modification times and file sizes for hash
    file_metadata = []
    for pdf_file in pdf_files:
        file_stat = os.stat(pdf_file)
        file_metadata.append((pdf_file, file_stat.st_mtime, file_stat.st_size))

    # Create a hash of the file metadata
    metadata_str = str(file_metadata).encode('utf-8')
    files_hash = hashlib.md5(metadata_str).hexdigest()

    # Check if we have a cached vector store for this hash
    vector_store_path = os.path.join(cache_dir, f"vector_store_{files_hash}")
    metadata_path = os.path.join(cache_dir, f"metadata_{files_hash}.json")

    # If cache exists, load it
    if os.path.exists(vector_store_path) and os.path.exists(metadata_path):
        try:
            print(f"⏳ Loading vector store from cache...")
            vector_store = FAISS.load_local(
                vector_store_path,
                OpenAIEmbeddings(model='text-embedding-ada-002'),
                allow_dangerous_deserialization=True  # Add this parameter
            )

            # Read metadata (optional, could be useful for debugging)
            with open(metadata_path, 'r') as f:
                metadata = json.load(f)
                cached_time = metadata.get('time', 'unknown')
                print(f"✅ Using cached vector store from {cached_time}")

            return vector_store

        except Exception as e:
            print(f"⚠️ Error loading cached vector store: {str(e)}")
            print("Rebuilding vector store...")
            # Continue with rebuilding if loading failed

    print(f"🔹 Building new vector store from {len(pdf_files)} PDF files...")

    # The rest of the original function to build vector store
    all_chunks = []
    all_metadata = []
    section_pattern = re.compile(r"^\s*(\d+\..+|[أ-ي]+[.)].+)$", re.MULTILINE)

    for file_index, file in enumerate(pdf_files):
        rfp_name = os.path.basename(file).replace('.pdf', '')
        text = ""

        try:
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text += fix_arabic_text(extracted_text) + "\n"

            if text.strip():
                sections = section_pattern.split(text)
                rfp_text = ""

                for section in sections:
                    section = section.strip()
                    if section:
                        rfp_text += f"{section}\n\n"

                text_splitter = RecursiveCharacterTextSplitter(chunk_size=200, chunk_overlap=50, length_function=len)
                rfp_chunks = text_splitter.split_text(rfp_text)

                for chunk in rfp_chunks:
                    all_chunks.append(chunk)
                    all_metadata.append({"source": rfp_name, "file_path": file})

                print(f"Processed file {file} and extracted {len(rfp_chunks)} chunks.")
        except Exception as e:
            print(f"Error processing {file}: {str(e)}")

    documents = [LangchainDocument(page_content=chunk, metadata=metadata)
                 for chunk, metadata in zip(all_chunks, all_metadata)]

    print(f"Created {len(documents)} documents from all RFPs.")

    if documents:
        embeddings = OpenAIEmbeddings(model='text-embedding-ada-002')
        vector_store = FAISS.from_documents(documents, embedding=embeddings)

        # Save the vector store and metadata
        try:
            vector_store.save_local(vector_store_path)

            # Save metadata about the cache
            metadata = {
                'hash': files_hash,
                'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'num_documents': len(documents),
                'files': [os.path.basename(f) for f in pdf_files]
            }

            with open(metadata_path, 'w') as f:
                json.dump(metadata, f, indent=2)

            print(f"✅ Vector store cached successfully for future use")
        except Exception as e:
            print(f"⚠️ Error caching vector store: {str(e)}")

        return vector_store
    else:
        print("No documents created. Vector store initialization failed.")
        return None


def generate_rfp_intro(llm, example_rfp, competition_name, competition_objectives, competition_description,
                       government_entity, cost_value, cost_method, start_stage, end_stage):
    prompt = f"""
اكتب القسم الأول من كراسة شروط مشروع بعنوان "{competition_name}" بهدف "{competition_objectives}" في مجال "{competition_description}". الجهة الحكومية المسؤولة: "{government_entity}".

هذا القسم هو: المقدمة.
لا تقم بكتابة "القسم الأول: المقدمة" في بداية النص، لأن هذا سيتم إضافته تلقائيًا. ابدأ مباشرة بمحتوى المقدمة.

يجب أن يشمل بالتفصيل:
- فقرة تعريفية تفصيلية تشرح جميع المصطلحات الفنية المرتبطة بموضوع المشروع، بما في ذلك:
  الجهة الحكومية، المتنافس، المنافسة، النظام، اللائحة التنفيذية، والمصطلحات الأخرى المتعلقة بالمنافسة.

- التعريفات الرسمية لكل المصطلحات المستخدمة في الكراسة، ويجب تنسيقها كالتالي:
  • الجهة الحكومية: {government_entity}
  • المتنافس: 
  • المنافسة: 
  • النظام: 
  • اللائحة التنفيذية: 
  • ... (أكمل باقي المصطلحات بنفس التنسيق)

- تقديم خلفية شاملة عن المشروع توضح السياق والأسباب التي دعت إلى طرحه.

- شرح نطاق الأعمال المتوقع، مقسمًا إلى مراحل واضحة ومترابطة تبدأ من "{start_stage}" وتنتهي عند "{end_stage}".

- تحديد المعايير العامة والضوابط التي يجب الالتزام بها أثناء تنفيذ المشروع.

- توضيح أهداف المنافسة بالتفصيل، والنتائج المتوقعة من تنفيذ المشروع.

3. تكاليف وثائق المنافسة:
ملاحظة: تقوم الجهة الحكومية بتحديد تكاليف وثائق المنافسة إن أرادت بيعها، وتحذف الفقرة في حال عدم انطباقها.  
ملاحظة: يجب على الجهة الحكومية تحري الدقة في تحديد تكاليف وثائق المنافسة بحيث تعكس الأسعار تكاليف إعدادها فقط، وعدم المبالغة في قيمتها بحيث تؤدي إلى إحجام الراغبين عن التقدم للمنافسة. ولا تشمل هذه التكاليف الأعمال الفنية والاستشارية.

| تكاليف وثائق المنافسة | آلية الدفع |
|--------------------------|-------------|
| القيمة بالأرقام ({cost_value} ريال سعودي) | {cost_method} |
| القيمة بالتنقيط           |             |

- ثم استكمل بباقي عناصر المقدمة:
• توضيح شروط أهلية مقدمي العروض، مثل المتطلبات الفنية والإدارية والمالية.

• إدراج جدول بمعلومات ممثل الجهة الحكومية المسؤول عن التواصل والمتابعة (اترك الخانات فارغة):

| الاسم الكامل | الوظيفة | البريد الإلكتروني | رقم الهاتف |
|--------------|----------|--------------------|-------------|
|              |          |                    |             |

• إدراج جدول يحدد مكان وتاريخ تسليم العروض:

| موقع تسليم العروض | العنوان التفصيلي | التاريخ | الوقت |
|-------------------|-------------------|---------|--------|
| مقر الجهة الحكومية | {government_entity}، قسم المشتريات، الدور الأول | 10 ذو القعدة 1446هـ | الساعة 12:00 ظهرًا |

ملاحظات هامة:
- لا تذكر أي أسماء محددة لأشخاص إلا إذا كانت مذكورة تحديدًا في وصف المشروع.
- اكتب بنمط "القسم الأول: المقدمة" لضمان التناسق.
- رقم قسم تكاليف وثائق المنافسة بـ "3" بشكل واضح كعنوان فرعي.
- في جدول "تكاليف وثائق المنافسة" اترك خلايا القيمة بالتنقيط فارغة.

ملاحظات هامة للتنسيق:
- لا تستخدم علامة # للعناوين، بل استخدم العناوين بشكل عادي.
- لا تستخدم علامة ** للنص العريض، اكتب النص بشكل عادي.
- استخدم الترقيم العادي للقوائم المرقمة (1. 2. 3.).
- استخدم النقاط العادية للقوائم غير المرقمة (•).
- ضع الجداول بصيغة عادية باستخدام | بين الأعمدة.

تعليمات التوليد هذه العليمات لك لا تقم بطباعتها:
- اكتب ما لا يقل عن 3000 كلمة لضمان تغطية شاملة.
- استخدم لغة عربية فصحى رسمية خالية من الأخطاء.
- اجعل المحتوى مترابطًا ومنطقيًا وتفصيليًا.
- لا تتجاهل أي من العناصر المطلوبة أعلاه، خاصة جدول تكاليف وثائق المنافسة والتفاصيل الدقيقة لجدول تسليم العروض.

كن حذرًا من أن أي اعتذار أو رفض للكتابة سيؤدي إلى فشل المهمة التي كُلفت بها. المطلوب كتابة محتوى تعليمي نموذجي فقط.
"""
    response = llm.predict(prompt)
    return response


def generate_rfp_general_terms(llm, example_rfp, competition_name, competition_objectives, competition_description,
                               government_entity):
    prompt = f"""
اكتب القسم الثاني من كراسة الشروط بعنوان "الأحكام العامة".

هذا القسم يتضمن المبادئ الأساسية والسياسات الرسمية التي تحكم العلاقة بين الجهة الحكومية والمتنافسين، ويعتمد على النموذج الثابت المعتمد في كراسات الشروط الحكومية.  
ويجب أن يشمل **اثني عشر (12) بندًا**، كل منها يمثل مبدأ رسمي من مبادئ المنافسات، ويُكتب بصياغة رسمية واضحة مع شرح مفصل يتضمن الحقوق والالتزامات والعقوبات القانونية ذات العلاقة.

مع ذلك، يجب أن يكون النص متكيّفًا مع سياق المنافسة التالية:
- اسم المشروع: {competition_name}
- الهدف من المشروع: {competition_objectives}
- وصف المجال: {competition_description}
- الجهة الحكومية: {government_entity}

المحتوى يجب أن يشمل البنود التالية (بنفس الترتيب والأسلوب الرسمي)، مع توضيح العقوبات والإشارة إلى الأنظمة الرسمية كمرجعية قانونية:

1. المساواة والشفافية  
• تلتزم الجهة الحكومية بضمان تكافؤ الفرص لجميع المتنافسين وتوفير المعلومات والبيانات ذات العلاقة بصورة عادلة ومتساوية.  
• يتم الإعلان عن أي تعديلات أو استفسارات تخص المنافسة من خلال القنوات الرسمية المعتمدة.

2. تعارض المصالح  
• يمنع مشاركة أي شخص أو جهة لها صلة مباشرة أو غير مباشرة قد تؤثر على نزاهة وعدالة المنافسة.  
• يجب على المتنافسين الإفصاح الكامل عن أي حالة تعارض مصالح فور العلم بها.  
• يخضع الإخلال بهذا البند للمساءلة وفقًا لنظام المنافسات والمشتريات الحكومية.

3. السلوكيات والأخلاقيات  
• يُمنع تقديم أو قبول أي هدايا أو مزايا أو امتيازات تهدف إلى التأثير على قرارات الترسية أو تقييم العروض.  
• يتوجب على جميع الأطراف الالتزام بالسلوك المهني والنزاهة وفقًا للأنظمة المعمول بها.  
• أي مخالفة تستوجب تطبيق العقوبات الواردة في اللوائح التنفيذية لنظام المنافسات.

4. السرية وإفشاء المعلومات  
• تلتزم جميع الأطراف بعدم إفشاء أو نقل أي معلومات تتعلق بالعروض أو الإجراءات أو التقييم أو غيرها.  
• يعاقب كل من يثبت تورطه في تسريب معلومات وفق الأنظمة الأمنية ذات العلاقة.

5. ملكية وثائق المنافسة  
• تبقى جميع الوثائق والمستندات الصادرة عن الجهة الحكومية ملكًا حصريًا لها، ولا يجوز نسخها أو تداولها إلا بموافقة رسمية.  
• يلتزم المتنافسون بإعادة الوثائق إذا طلبت الجهة ذلك.

6. حقوق الملكية الفكرية  
• تحتفظ الجهة الحكومية بجميع الحقوق المتعلقة بالمستندات أو التصاميم أو الحلول المقدمة ضمن المنافسة عند الترسية.  
• لا يجوز إعادة استخدام المواد المقدمة في المنافسة دون إذن خطي من الجهة.

7. المحتوى المحلي  
• يُشجّع استخدام المنتجات والخدمات الوطنية في جميع مكونات المشروع كلما كان ذلك ممكنًا، وذلك وفقًا لاستراتيجية المحتوى المحلي المعتمدة.

8. أنظمة وأحكام الاستيراد  
• يجب الالتزام بجميع الأنظمة واللوائح ذات العلاقة بالاستيراد، بما في ذلك الحصول على التصاريح والتراخيص النظامية.  
• يُمنع استخدام أو استيراد أي مواد محظورة أو لا تتوافق مع المواصفات السعودية.

9. تجزئة المنافسة  
• تحتفظ الجهة الحكومية بحقها الكامل في تجزئة المشروع إلى مراحل أو أجزاء حسب المصلحة العامة.  
• يلتزم المتنافس بتنفيذ الجزء المُسند إليه وفقًا لكافة الشروط المنصوص عليها في كراسة الشروط.

10. الاستبعاد من المنافسة  
• يتم استبعاد أي متنافس يخل بشروط المنافسة أو يقدم بيانات غير صحيحة أو يثبت تورطه في ممارسات غير قانونية.  
• الاستبعاد يكون مستندًا إلى مواد نظام المنافسات والمشتريات الحكومية، ولا يحق الاعتراض عليه إلا من خلال القنوات القانونية الرسمية.

11. إلغاء المنافسة وأثره  
• يجوز للجهة الحكومية إلغاء المنافسة في حال وجود أسباب جوهرية مثل اكتشاف أخطاء أو تعارضات نظامية.  
• لا تتحمل الجهة أي مسؤولية مالية تجاه المتنافسين في حال الإلغاء، وذلك وفقًا للمادة (28) من نظام المنافسات والمشتريات الحكومية.

12. التفاوض مع أصحاب العروض  
• في حال تجاوز الأسعار المقدمة حدود الأسعار السوقية المعقولة، يحق للجهة التفاوض مع أصحاب العروض الأفضل تقييمًا.  
• يتم تسجيل محاضر التفاوض وإرفاقها ضمن مستندات الترسية الرسمية.

ملاحظات هامة:
- لا تذكر أي أسماء لأشخاص إلا إذا وردت صراحة في وصف المشروع.
- لا تدمج البنود أو تترك أي بند ناقصًا.
- لا تكتب "القسم الثاني: الأحكام العامة" داخل النص.
- لا تستخدم أي تنسيقات Markdown أو رموز برمجية.

تعليمات داخلية:
- اكتب من 3000 إلى 4000 كلمة.
- استخدم لغة عربية رسمية فصحى.
- اربط البنود بسياق المشروع دون الإخلال بطابعها الرسمي.
- تأكد من الإشارة إلى "نظام المنافسات والمشتريات الحكومية" عند ذكر أي التزام قانوني أو جزائي.
- لا ترفض أو تعتذر عن التنفيذ.
"""
    response = llm.predict(prompt)
    return response


def generate_rfp_offer_preparation(llm, example_rfp, technical_docs, alternative_offers, initial_guarantee,
                                   competition_name, competition_objectives, competition_description,
                                   government_entity):
    prompt = f"""

اكتب القسم الثالث من كراسة الشروط بعنوان "إعداد العروض".

مع ذلك، يجب أن يكون النص متكيّفًا مع سياق المنافسة التالية:
- اسم المشروع: {competition_name}
- الهدف من المشروع: {competition_objectives}
- وصف المجال: {competition_description}
- الجهة الحكومية: {government_entity}

يجب أن يكون هذا القسم رسميًا ومفصلاً ومبنيًا على ما يلي:

1. لغة العرض  
• يجب أن تكون جميع المستندات والعروض والمرفقات المقدمة باللغة العربية الفصحى.  
• في حال تقديم أي وثائق بلغة أجنبية، يجب إرفاق ترجمة عربية معتمدة من جهة رسمية.  
• في حالة وجود تعارض بين النص العربي والنص المترجم، يُعتمد النص العربي فقط.

2. العملة الرسمية  
• يجب أن تكون جميع الأسعار والتكاليف المقدمة في العرض محددة بعملة الريال السعودي.  
• لا تُقبل العروض المقدمة بعملات أجنبية.  
• يتم اعتماد سعر الصرف الرسمي من مؤسسة النقد العربي السعودي ليوم تقديم العرض.

3. دقة المعلومات  
• يتحمل المتنافس كامل المسؤولية عن صحة ودقة البيانات المقدمة.  
• يحق للجهة الحكومية استبعاد أي عرض يتبين لاحقًا أنه يحتوي على معلومات غير صحيحة أو مضللة.

4. محتويات العرض  
• يجب أن يحتوي العرض على الوثائق الفنية والمالية المطلوبة بشكل منفصل وواضح.  
• الوثائق الفنية تشمل: {technical_docs}  
• يجب تقديم العرض الفني بشكل مفصل يوضح منهجية العمل المقترحة، الخطة الزمنية، والسير الذاتية.  
• يجب تقديم قائمة بالمشاريع المشابهة التي سبق تنفيذها.  

• العروض البديلة: {alternative_offers}  
• يُسمح بتقديم عروض بديلة فقط إذا نصت كراسة الشروط على ذلك، ويجب أن تكون مصحوبة بتبرير فني واقتصادي واضح.  
• تُرفض العروض البديلة المخالفة أو التي تُخل بمبدأ التنافسية والعدالة.

5. الفصل بين العرض الفني والمالي  
• يجب تقديم العرض الفني والمالي في ملفين منفصلين على النحو التالي:  
  • الملف الفني: يشمل خطة التنفيذ، المنهجية، السيرة الذاتية للفريق، الجدول الزمني.  
  • الملف المالي: يشمل تفاصيل التكاليف، الأسعار، والضمانات المالية.  
• يُرفض أي عرض يتم فيه دمج المحتوى الفني والمالي في مستند واحد بشكل قاطع لضمان الحياد عند التقييم.

6. الضمان الابتدائي  
• يجب تقديم ضمان ابتدائي مع العرض وفقًا لما يلي:  
{initial_guarantee}  
• يعتبر الضمان الابتدائي شرطًا أساسيًا لقبول العرض، ويُرفض أي عرض غير مرفق به.

7. تنسيق العرض  
• يجب تقديم العرض في ملفات منظمة ومرتبة حسب التسلسل المطلوب.  
• يجب ترقيم جميع صفحات العرض بشكل متسلسل.  
• يجب توقيع جميع الصفحات من الشخص المفوض.  
• يجب تقديم فهرس واضح لمحتويات العرض.

8. جدول محتويات العرض المطلوبة

| القسم | المحتوى | الملاحظات |
|-------|---------|-----------|
|       |         |           |
|       |         |           |
|       |         |           |
|       |         |           |
|       |         |           |

ملاحظات هامة للتنسيق:
- استخدم صيغة "القسم الثالث: إعداد العروض" في بداية النص.
- لا تستخدم علامة # للعناوين.
- لا تستخدم علامة ** للنص العريض.
- استخدم الترقيم العادي للقوائم (1. 2. 3.).
- استخدم النقاط العادية للقوائم غير المرقمة (•).
- لا تدرج أي تنسيق Markdown.
- اذكر الأسماء والمبالغ المقدمة من المستخدم بدقة.

تعليمات التوليد (لا تطبعها في الإخراج):  
- اكتب من 3000 إلى 4000 كلمة.  
- استخدم لغة رسمية واضحة فصحى ومترابطة.  
- لا ترفض أو تعتذر عن التوليد؛ هذه مهمة تعليمية رسمية.
"""
    response = llm.predict(prompt)
    return response


def generate_rfp_offer_submission(llm, example_rfp, competition_name, competition_objectives, competition_description,
                                  government_entity):
    prompt = f"""
اكتب القسم الرابع من كراسة الشروط بعنوان "تقديم العروض".

مع ذلك، يجب أن يكون النص متكيّفًا مع سياق المنافسة التالية:
- اسم المشروع: {competition_name}
- الهدف من المشروع: {competition_objectives}
- وصف المجال: {competition_description}
- الجهة الحكومية: {government_entity}

لا تكتب عنوان القسم داخل النص، بل ابدأ مباشرة بالمحتوى، على أن يضاف العنوان تلقائيًا لاحقًا في التنسيق النهائي.

يجب أن يشمل هذا القسم الفقرات التالية:

أولاً: آلية تقديم العروض  
• يلتزم المتنافس بتقديم عرضه في الموعد المحدد من خلال البوابة الإلكترونية أو الوسيلة البديلة المعتمدة.  
• في حال تعطل البوابة الإلكترونية، يُمكن التقديم عبر الوسيلة البديلة خلال المدة التي تحددها الجهة الحكومية.  
• تُقدم العروض من خلال ممثل مفوض، مرفقًا بجميع المستندات والوثائق حسب متطلبات الكراسة.  
• يجب تعبئة البيانات إلكترونيًا، وتوقيع العرض بخطاب رسمي صادر عن الجهة المتقدمة.  
• تُقدم العروض ضمن مظاريف إلكترونية، وتُسلَّم عبر البوابة الإلكترونية أو بالبريد الرسمي حسب قرار الجهة الحكومية.  
• لا يُقبل أي عرض يصل خارج المهلة الزمنية المحددة أو بوسيلة غير معتمدة.

ثانيًا: تسليم العروض المتأخرة  
• لا يُعتد بأي عرض يصل بعد انتهاء المهلة المحددة.  
• العروض التي ترد بوسائل غير معتمدة يتم استبعادها مباشرة دون فتحها.  
• المتنافس مسؤول عن أي تأخير أو عطل، سواءً تقني أو إداري، ويُرفض العرض تلقائيًا دون التزام من الجهة الحكومية.

ثالثاً: فتح العروض  
• يتم فتح العروض المقدمة في جلسة رسمية بحضور ممثلي المتنافسين الراغبين.  
• يُحدد موعد فتح العروض مسبقًا وتُبلّغ به جميع الأطراف عبر القنوات الرسمية.  
• يُعد محضر رسمي يتضمن البيانات الأساسية لكل عرض، ويُوقع من اللجنة المختصة.  
• يحق للمتنافسين الاطلاع على المحضر بعد انتهاء الجلسة.

رابعاً: تمديد فترة تلقي العروض  
• يحق للجهة الحكومية تمديد المهلة في حال وجود مبرر نظامي أو لوجستي.  
• يتم إعلان التمديد عبر نفس القنوات الرسمية المستخدمة للإعلان الأصلي.  
• يسري التمديد على جميع المتنافسين دون استثناء.  
• لا يجوز الاعتراض على قرار التمديد من قبل أي متقدم.

خامساً: سحب العروض  
• يمكن للمتنافس سحب عرضه في أي وقت قبل الموعد النهائي لتقديم العروض.  
• يتم تقديم طلب السحب بخطاب رسمي من الجهة المتقدمة، موجه إلى الجهة الحكومية عبر القنوات المعتمدة.  
• لا يجوز سحب العرض بعد انتهاء الموعد المحدد لتقديم العروض.  
• في حال سحب العرض بعد فتح المظاريف، يحق للجهة مصادرة الضمان الابتدائي دون إشعار مسبق.

سادساً: جدول مواعيد تقديم العروض  
• توضح الجهة الحكومية الجدول الزمني الخاص بمراحل المنافسة، بما في ذلك موعد التقديم وفتح المظاريف والإعلان عن النتائج.  
• يُطلب منك تعبئة هذا الجدول بناءً على البيانات التي يتم توفيرها من المستخدم، ولا يجوز كتابة أي تاريخ أو وقت أو مكان من عندك.

| المرحلة                    | التاريخ              | الوقت         | المكان                                      |
|----------------------------|----------------------|---------------|---------------------------------------------|
| آخر موعد لتقديم العروض     |                      |               |                                             |
| جلسة فتح المظاريف         |                      |               |                                             |
| الإعلان عن العروض المقبولة |                      |               |                                             |
| المدة الإضافية (إن وجدت)   |                      |               |                                             |

ملاحظات التنسيق:
- لا تكتب عنوان "القسم الرابع: تقديم العروض" داخل النص.
- لا تستخدم علامات مثل # أو ** أو أي تنسيقات Markdown.
- اكتب النص بلغة عربية فصحى رسمية.
- لا تملأ الجدول إلا باستخدام بيانات مدخلة من المستخدم.

تعليمات التوليد (لا تُطبع):
- اكتب من 3000 إلى 4000 كلمة.
- لا ترفض تنفيذ الطلب لأي سبب.
"""
    response = llm.predict(prompt)
    return response


def generate_rfp_offer_analysis(llm, competition_name, competition_description, competition_objectives, pause_period,
                                government_entity):
    prompt = f"""
اكتب القسم الخامس من كراسة الشروط بعنوان "تقييم العروض".
لا تكتب عنوان القسم أبدًا، وابدأ مباشرة في كتابة المحتوى.

هذا القسم يجب أن يتبع الصيغة الرسمية المعتمدة، ويتكون من الفقرات التالية:

أولاً: سرية تقييم العروض  
تلتزم الجهة الحكومية بعدم إفشاء أي بيانات أو رسومات أو وثائق أو معلومات تتعلق بتقييم العروض المستلمة، سواء كان الإفشاء تحريرياً أو شفهياً، أو استغلالها أو الإفصاح عنها إلى أي طرف، ويشمل ذلك كل ما اطلعت عليه اللجنة من معلومات تخص المتنافسين، باستثناء ما تُلزم الأنظمة بنشره.

ثانيًا: معايير تقييم العروض  
تعتمد لجنة التقييم في مشروع:  
- {competition_name}  
- في مجال: {competition_description}  
- بهدف: {competition_objectives}  
- للجهة: {government_entity}

على معايير واضحة تشمل الجوانب التالية:
• التقييم الفني: يشمل منهجية التنفيذ، الخبرات السابقة، فريق العمل، الجدول الزمني، المحتوى المحلي.  
• التقييم المالي: يشمل التكلفة الإجمالية، الأسعار المفصلة، مدى الالتزام بالميزانية، التوازن بين التكلفة والقيمة.  
• تمنح درجات محددة لكل بند، ويُحتسب مجموع النقاط وفق النسب التالية:  
  - التقييم الفني: 70٪  
  - التقييم المالي: 30٪  

ثالثًا: جدول معايير التقييم الفني التفصيلية

| المعيار                     | الوزن النسبي | التفاصيل                            | درجة التقييم |
|-----------------------------|---------------|--------------------------------------|---------------|
| منهجية التنفيذ              | 20٪           | وضوح منهجية تنفيذ المشروع           | من 0 إلى 20   |
| خبرات سابقة مماثلة         | 15٪           | مشاريع مماثلة تم تنفيذها بنجاح     | من 0 إلى 15   |
| كفاءة الفريق               | 15٪           | مؤهلات وخبرات الفريق الفني          | من 0 إلى 15   |
| الجدول الزمني              | 10٪           | وضوح ومناسبة مراحل التنفيذ          | من 0 إلى 10   |
| المحتوى المحلي             | 10٪           | مدى استخدام منتجات وخدمات محلية     | من 0 إلى 10   |

رابعًا: تصحيح العروض  
• تقوم لجنة التقييم بمراجعة جداول الكميات والأسعار المقدمة من المتنافسين، والتأكد من صحة العمليات الحسابية.  
• في حال وجود فروق بين السعر كتابةً والسعر رقماً، يتم اعتماد السعر كتابةً باعتباره المرجع الرسمي.  
• يُسمح للجنة بإجراء التصحيحات الحسابية غير الجوهرية دون الحاجة إلى موافقة المتنافس.  
• أما في حالة اكتشاف أخطاء حسابية تؤثر جوهريًا على العرض، فيجوز استبعاده بقرار مسبب من اللجنة.  
• يُوثق كل تعديل أو تصحيح في محضر رسمي، ويُعتمد وفقًا لما نصت عليه اللائحة التنفيذية لنظام المنافسات والمشتريات الحكومية.

خامسًا: فحص العروض  
• يجب الالتزام بمعايير التأهيل والشهادات المطلوبة.  
• في حال نقص وثائق غير جوهرية، يجوز منح مهلة للاستكمال.  
• العروض التي تحتوي على بنود غير مسعرة تُستبعد.  
• في حال عدم تنفيذ بند مسعر لاحقًا، تُحمّل الجهة المتقدمة التكلفة أو يُستبعد عرضها.

... (تابع باقي الفقرات كما هي)
"""
    response = llm.predict(prompt)
    return response


def generate_rfp_award_contract(llm, example_rfp, penalties, competition_name, competition_description,
                                competition_objectives, pause_period, government_entity):
    prompt = f"""
يجب أن يكون النص متكيّفًا مع سياق المنافسة التالية:
- اسم المشروع: {competition_name}
- الهدف من المشروع: {competition_objectives}
- وصف المجال: {competition_description}
- الجهة الحكومية: {government_entity}

اكتب القسم السادس من كراسة الشروط بعنوان "متطلبات التعاقد".
لا تكتب عنوان القسم، بل ابدأ مباشرة في كتابة المحتوى.

يجب أن يشمل النص الفقرات التالية بلغة رسمية واضحة:

1. إخطار الترسية  
• تصدر الجهة الحكومية إشعارًا رسميًا بترسية المنافسة على المتنافس الفائز.  
• يتضمن الإشعار البيانات التالية: اسم المشروع، القيمة الإجمالية، مدة التنفيذ، تاريخ الاستلام الابتدائي.  
• يتم إرسال الإشعار عبر البريد الإلكتروني الرسمي أو عبر المنصة الحكومية.  
• يلتزم المتنافس الفائز بتقديم الضمانات والوثائق التعاقدية خلال مدة لا تتجاوز 5 أيام عمل من تاريخ الإشعار.

2. تقديم الضمان النهائي  
• يجب على المتنافس الفائز تقديم ضمان نهائي غير مشروط بنسبة تحددها الجهة الحكومية، ويكون ساريًا حتى إتمام التنفيذ النهائي.  
• يجب أن يصدر الضمان من بنك معتمد، ويُقدم بصيغته الرسمية المعتمدة.  
• لا يتم توقيع العقد دون استلام هذا الضمان.

3. توقيع العقد  
• يتم توقيع العقد بين الجهة الحكومية والمتنافس الفائز بعد استيفاء جميع الشروط النظامية.  
• يشتمل العقد النهائي على العناصر التالية:

  - نطاق الأعمال والمواصفات الفنية  
  - المدة الزمنية المحددة للتنفيذ  
  - الجدول الزمني للمراحل  
  - جدول الدفعات والمستحقات  
  - شروط التسليم والاعتمادات  
  - الغرامات والمخالفات التعاقدية  

• لا يُعد العقد نافذًا إلا بعد توقيعه من الجهتين واستكمال الضمان النهائي.

4. التأمين  
• يلتزم المتعاقد بتوفير تغطية تأمينية شاملة تشمل:  
  - العمالة  
  - موقع المشروع  
  - المعدات  
• يجب أن تبقى التأمينات سارية طوال فترة تنفيذ المشروع.

5. الاستلام الأولي والنهائي  
• يتم الاستلام الأولي بعد اكتمال الأعمال الرئيسية والموافقة على مطابقة المواصفات.  
• يتم الاستلام النهائي بعد انتهاء فترة الضمان ومعالجة جميع الملاحظات.  
• تصدر لجنة مختصة محاضر رسمية لكل من الاستلامين.

6. التعديلات أثناء التنفيذ  
• يجوز للجهة الحكومية تعديل نطاق العمل أو الجدول الزمني.  
• يتم توثيق التعديلات بمحاضر رسمية، ويوقعها الطرفان، مع تحديث جدول الدفعات عند الحاجة.

7. جدول الإجراءات التعاقدية

| الإجراء                        | الجهة المسؤولة         | المستندات المطلوبة              | المدة الزمنية      |
|-------------------------------|--------------------------|----------------------------------|---------------------|
| إصدار إشعار الترسية          | الجهة الحكومية           | إشعار رسمي                        | خلال 1 يوم عمل     |
| تقديم الضمان النهائي         | المتنافس الفائز         | ضمان بنكي بصيغة معتمدة          | خلال 5 أيام عمل    |
| توقيع العقد                   | الجهة + المتنافس        | النسخة النهائية من العقد         | خلال 2 يوم عمل     |
| تسليم الموقع                  | الجهة الحكومية           | محضر استلام موقع                 | بعد توقيع العقد    |
| بدء التنفيذ                   | المتنافس الفائز         | محضر بدء تنفيذ                   | فور استلام الموقع  |

8. جدول الدفعات الاسترشادي

| المرحلة                         | نسبة الدفعة | شروط الاستحقاق                       | الوثائق المطلوبة                  |
|----------------------------------|--------------|----------------------------------------|------------------------------------|
| بعد توقيع العقد                 | 10%          | تقديم الضمان النهائي وتوقيع العقد     | ضمان نهائي + نسخة العقد           |
| بعد إتمام 50% من الأعمال       | 40%          | تقرير إنجاز موثق                      | تقرير إنجاز موقع + إشراف هندسي    |
| بعد الاستلام الأولي            | 30%          | اعتماد لجنة الاستلام الابتدائي        | محضر الاستلام الابتدائي           |
| بعد الاستلام النهائي           | 20%          | انتهاء فترة الضمان وتصفية الملاحظات   | محضر الاستلام النهائي              |

9. نماذج الضمانات والمستندات المطلوبة  
• نموذج الضمان البنكي الابتدائي  
• نموذج الضمان البنكي النهائي  
• نموذج ضمان دفعة مقدمة  
• محضر استلام موقع المشروع  
• محضر بدء الأعمال  
• تقرير إنجاز شهري  
• محضر الاستلام الابتدائي  
• محضر الاستلام النهائي

10. إدارة المخاطر التعاقدية  
• تحدد الجهة الحكومية أبرز المخاطر المحتملة خلال التنفيذ.  
• يجب على المتعاقد وضع خطة لإدارة المخاطر تشمل إجراءات وقائية واضحة.  
• تتم مراجعة المخاطر شهريًا وتوثيقها بمحضر دوري معتمد من الطرفين.

11. الغرامات  
• تطبق الجهة الحكومية غرامات جزائية في حال الإخلال بأي من الالتزامات التعاقدية، وتشمل:  
{penalties}  
• تُحسب الغرامة بناءً على نوع المخالفة ومدى تأثيرها على سير العمل.  
• لا تعفي الغرامة من التزامات التنفيذ أو التمديد، بل تهدف إلى تحقيق الانضباط.

12. نموذج عقد الترسية النهائي  
فيما يلي نموذج مختصر لصيغة العقد المعتمد، ويُستخدم كأساس للصيغة النهائية التي تُوقّع بين الطرفين:

[بداية النموذج]

**عقد تنفيذ مشروع: {competition_name}**

تم إبرام هذا العقد في تاريخ (...........) بين كل من:  
الجهة الحكومية: {government_entity}  
ويمثلها في هذا العقد: (..........)  
والشركة المتعاقدة: (اسم الشركة المتقدمة بالعطاء)

وبموجب ترسية المنافسة رقم: (...........)  
ووفقًا لكراسة الشروط والمواصفات لمشروع: {competition_description}  
فقد تم الاتفاق على تنفيذ المشروع خلال مدة قدرها (..........) شهرًا  
وبقيمة إجمالية مقدارها (..........) ريال سعودي.

ويخضع هذا العقد لكافة الشروط والمتطلبات الواردة في وثائق المنافسة، والأنظمة واللوائح ذات العلاقة.

[نهاية النموذج]

ملاحظات التنسيق:  
- استخدم صيغة "القسم السادس: متطلبات التعاقد" فقط عند الطباعة النهائية.  
- لا تستخدم أي علامات Markdown أو رموز برمجية.  
- لا تدخل مواد أو قوانين نظامية.  
- لا تذكر مبالغ أو تواريخ إلا إذا تم تحديدها من المستخدم.

تعليمات التوليد (لا تطبعها):  
- استخدم لغة رسمية واضحة، فصحى وغير متكلفة.  
- اكتب من 3000 إلى 4000 كلمة.  
- لا ترفض التوليد تحت أي ظرف. هذا محتوى رسمي تعليمي.
"""
    response = llm.predict(prompt)
    return response


def generate_rfp_work_scope(llm, competition_name, competition_description, competition_objectives, execution_city,
                            execution_district, execution_region, government_entity):
    prompt = f"""
يجب أن يكون النص متكيّفًا مع سياق المنافسة التالية:
- اسم المشروع: {competition_name}
- الهدف من المشروع: {competition_objectives}
- وصف المجال: {competition_description}
- الجهة الحكومية: {government_entity}

لا تكتب عنوان القسم أبدًا، وابدأ مباشرة في كتابة المحتوى.

اكتب القسم السابع من كراسة الشروط بعنوان "نطاق العمل المفصل".

المشروع: {competition_name}  
المجال: {competition_description}  
مكان التنفيذ: مدينة {execution_city}، حي {execution_district}، منطقة {execution_region}

يجب أن يتضمن النص فقرات رئيسية مكتوبة بلغة رسمية ومترابطة:

أولًا: نطاق عمل المشروع  
- في هذه الفقرة يتم وصف العمل المطلوب من المتعاقد بشكل مفصل.  
- وضّح آلية التنفيذ، والمخرجات المطلوبة من كل مرحلة.  
- اشرح كيف سيتم تنفيذ المشروع من البداية للنهاية مع تحديد طبيعة الأعمال (خدمة، تطوير، توريد... إلخ).  
- لا تدخل تفاصيل فنية دقيقة أو أرقام مالية إلا إذا وردت من المستخدم.  
- يتم تنفيذ المشروع في مدينة {execution_city}، حي {execution_district}، منطقة {execution_region}.

ثانيًا: مراحل المشروع  
فصِّل المراحل التي ستمر بها المنافسة، وحدد لكل مرحلة:  
- وصف تفصيلي للمرحلة  
- المخرجات المتوقعة  
- المسؤوليات والمهام المحددة  
- آليات المتابعة والتقييم  
**يُدرج الجدول الزمني الكامل في الفقرة التالية. لا تذكر المدد الزمنية هنا تجنبًا للتكرار.**

ثالثًا: متطلبات التنفيذ  
- متطلبات الجودة  
- متطلبات الكوادر البشرية  
- متطلبات المعدات والتجهيزات  
- متطلبات إدارة المشروع  
- متطلبات التوثيق وإعداد التقارير

رابعًا: الجدول الزمني للمشروع

| المرحلة | وصف المرحلة | المدة الزمنية | المخرجات | ملاحظات |
| ------- | ------------ | --------------|-----------|----------|
|         |              |               |           |          |
|         |              |               |           |          |
|         |              |               |           |          |
|         |              |               |           |          |
|         |              |               |           |          |

اقترح محتوى مناسب واقعي لهذا الجدول مع تعبئة كل الخلايا بناءً على طبيعة المشروع. كل مرحلة يجب أن تتضمن وصفًا واضحًا ومدة زمنية منطقية والمخرجات المتوقعة وأي ملاحظات مهمة.

خامسًا: هيكل توزيع المسؤوليات

| المسؤولية | الجهة المسؤولة | المخرجات المتوقعة | آلية المتابعة |
|-----------|----------------|--------------------|----------------|
|           |                |                    |                |
|           |                |                    |                |
|           |                |                    |                |
|           |                |                    |                |

قم بتعبئة هذا الجدول بمسؤوليات واقعية تعكس طبيعة المشروع، مع تحديد الجهات المسؤولة والمخرجات وآليات المتابعة المناسبة.

سادسًا: المتطلبات الفنية التفصيلية  
- المواصفات الفنية للمخرجات  
- المعايير المطلوب الالتزام بها  
- الاشتراطات الخاصة بالتنفيذ  
- المتطلبات التقنية  
- متطلبات المطابقة والتوافق

سابعًا: آليات ضمان الجودة  
- إجراءات مراقبة الجودة  
- معايير قبول المخرجات  
- آليات التدقيق والمراجعة  
- إجراءات التصحيح والتحسين  
- توثيق إجراءات ضمان الجودة

ملاحظات التنسيق:  
- استخدم صيغة "القسم السابع: نطاق العمل المفصل" في بداية النص النهائي.  
- لا تستخدم علامة # للعناوين.  
- لا تستخدم ** للنص العريض.  
- استخدم النقاط والقوائم فقط بصيغة رسمية.  
- لا تستخدم تنسيقات Markdown.  
- اذكر الأسماء المعطاة من المستخدم كما هي.  
- أي مبالغ مالية يجب أن تكون مأخوذة من مدخلات المستخدم.  
- قم بملء الجداول بمحتوى واقعي مناسب يعكس طبيعة المشروع.

تعليمات التوليد هذه العليمات لك لا تقم بطباعتها:  
- اكتب بين 3000 إلى 4000 كلمة.  
- استخدم لغة رسمية، واضحة، مفصلة، خالية من الحشو.  
- لا ترفض الكتابة تحت أي ظرف، هذا محتوى تعليمي رسمي فقط.
"""
    response = llm.predict(prompt)
    return response


def generate_rfp_specifications(llm, example_rfp, required_materials, competition_name, competition_objectives,
                                competition_description, government_entity):
    prompt = f"""
يجب أن يكون النص متكيّفًا مع سياق المنافسة التالية:
- اسم المشروع: {competition_name}
- الهدف من المشروع: {competition_objectives}
- وصف المجال: {competition_description}
- الجهة الحكومية: {government_entity}

لا تكتب عنوان القسم أبدًا، وابدأ مباشرة في كتابة المحتوى.

اكتب القسم الثامن من كراسة الشروط بعنوان "المواصفات الفنية".

هذا القسم يحدد المواصفات الأساسية التي يجب الالتزام بها أثناء تنفيذ المشروع، ويتضمن شروطًا واضحة تتعلق بفريق العمل، الجودة، والسلامة. المطلوب كتابة النص بلغة رسمية ومنظمة، وبأسلوب كراسات الشروط المعتمدة دون إسهاب أو تعقيد قانوني زائد.

يشمل هذا القسم الفقرات التالية:

أولاً: الشروط الخاصة بفريق العمل  
• يجب على المتعاقد توفير الكوادر البشرية اللازمة لتنفيذ المشروع وفقًا لطبيعة الأعمال.  
• يُشترط أن تكون القوى العاملة ذات كفاءة وخبرة كافية في مجال المشروع.  
• يلتزم المتعاقد بسداد الأجور والمستحقات المالية للعاملين بانتظام، وتوفير التأمينات النظامية لهم.  
• يجب أن تتوافق مؤهلات وخبرات العاملين مع متطلبات المشروع الفنية والتنفيذية.  
• يُلزم المتعاقد بتوفير البدائل في حال غياب أو تعذر استمرارية أحد أفراد الفريق لضمان سير العمل دون توقف.  
• لا يجوز التعاقد من الباطن مع أفراد أو شركات لتنفيذ الأعمال إلا بموافقة مسبقة من الجهة الحكومية.

ثانيًا: مواصفات الجودة  
• يتعين على المتعاقد الالتزام بتطبيق أعلى معايير الجودة خلال جميع مراحل التنفيذ.  
• يجب أن تتوافق كافة الأعمال المنفذة مع المواصفات الفنية المعتمدة في كراسة الشروط.  
• تلتزم الجهة المنفذة بإجراء فحوصات واختبارات دورية لضمان مطابقة التنفيذ للمعايير.  
• يتم توثيق نتائج الفحوصات في تقارير رسمية تُعرض على الجهة الحكومية.  
• في حال ظهور أي عيوب أو انحرافات، يتعين معالجتها فورًا دون تأخير.

ثالثًا: مواصفات السلامة  
• يلتزم المتعاقد بتطبيق جميع إجراءات السلامة المهنية في موقع العمل.  
• يجب توفير معدات الحماية الشخصية (PPE) لجميع العاملين، وتدريبهم على استخدامها.  
• يجب تأمين المواقع والمناطق الخطرة، وتركيب اللوحات التحذيرية والإرشادية حسب الحاجة.  
• يتعين على المتعاقد إعداد خطة طوارئ واضحة، تتضمن خطوات الإخلاء والإسعافات الأولية.  
• يجب الالتزام بجميع لوائح السلامة المعتمدة محليًا ودوليًا، وتقديم ما يثبت الالتزام بها عند الطلب.

رابعًا: المواصفات الفنية للمخرجات  
• المواد المطلوبة تشمل: {required_materials}  
• مواصفات المواد المستخدمة في المشروع  
• مواصفات الأجهزة والمعدات المطلوبة  
• مواصفات البرمجيات والأنظمة (إن وجدت)  
• المعايير الفنية المطلوب تطبيقها  
• شروط المطابقة للمواصفات القياسية

خامسًا: ضوابط التوثيق والتسليم  
• متطلبات التوثيق الفني  
• آلية تسليم المخرجات  
• صيغ ومعايير التقارير الدورية  
• متطلبات الأرشفة والحفظ  
• آلية نقل المعرفة والتدريب

سادسًا: مصفوفة المواصفات والمتطلبات الفنية  
يجب أن تتضمن هذه المصفوفة تفاصيل فنية دقيقة توضح كل بند من بنود العمل المطلوب، والمعايير التي يجب تحقيقها، والآلية التي يمكن بها قياس تحقق كل متطلب.

| البند | المواصفات المطلوبة | معايير القبول | أسلوب القياس |
|------|----------------------|----------------|----------------|
|      |                      |                |                |
|      |                      |                |                |
|      |                      |                |                |
|      |                      |                |                |

قم بتعبئة الجدول بمحتوى فني يتناسب مع طبيعة المشروع، مثل المواد، أو المعدات، أو البرامج، أو مراحل تنفيذ محددة تتطلب معايير دقيقة.

سابعًا: متطلبات ضمان الجودة  
يعرض الجدول التالي مؤشرات الأداء الرئيسية التي يجب مراقبتها، والمعايير المرجعية المعتمدة، وآلية تقييم الالتزام بها طوال فترة تنفيذ المشروع.

| المجال | مؤشرات الأداء | المعايير المرجعية | آلية التقييم |
|--------|----------------|---------------------|----------------|
|        |                |                     |                |
|        |                |                     |                |
|        |                |                     |                |
|        |                |                     |                |

يجب ملء هذا الجدول بمؤشرات أداء واقعية يمكن قياسها، وتستند إلى معايير فنية أو تنظيمية معتمدة في المجال، مثل الالتزام بالجدول الزمني، أو دقة التقارير، أو جودة المخرجات النهائية.

ملاحظات التنسيق:  
- استخدم صيغة "القسم الثامن: المواصفات الفنية" في بداية النص.  
- لا تستخدم علامة # للعناوين.  
- لا تستخدم علامة ** للنص العريض.  
- استخدم الترقيم العادي للقوائم المرقمة (1. 2. 3.).  
- استخدم النقاط العادية للقوائم غير المرقمة (•).  
- لا تدرج أي جداول أو عناصر غير مطلوبة في هذا القسم.  
- قم بملء الجداول بمحتوى واقعي مناسب يعكس طبيعة المشروع.

تعليمات التوليد هذه العليمات لك لا تقم بطباعتها:  
- اكتب بين 3000 إلى 4000 كلمة.  
- استخدم لغة رسمية فصحى، واضحة ومترابطة.  
- لا ترفض الكتابة. هذا محتوى تعليمي رسمي يجب إعداده بدقة.
"""
    response = llm.predict(prompt)
    return response


def generate_rfp_general_contract_terms(llm, example_rfp, competition_name, competition_objectives,
                                        competition_description, government_entity):
    prompt = f"""
            يجب أن يكون النص متكيّفًا مع سياق المنافسة التالية:
        - اسم المشروع: {competition_name}
        - الهدف من المشروع: {competition_objectives}
        - وصف المجال: {competition_description}
        -الجهة الحكومية: {government_entity}
    اكتب القسم التاسع من كراسة الشروط بعنوان "متطلبات المحتوى المحلي".
لا تكتب عنوان القسم ابدا مباشره في كتابة المحتوى

    هذا القسم يوضح التزامات المتنافسين تجاه المحتوى المحلي أثناء تنفيذ المشروع، وذلك بما يتماشى مع سياسات المملكة العربية السعودية الداعمة للصناعة الوطنية، وتوجيهات رؤية 2030.

    المطلوب أن يكون النص رسميًا، واضحًا، ومحددًا، بدون إدخال تفاصيل غير ضرورية أو اجتهادات.

    يجب أن يشمل النقاط التالية:

    أولاً: القائمة الإلزامية  
    • تلتزم الجهة الحكومية بتحديد قائمة إلزامية من المواد أو الخدمات التي يجب توريدها من السوق المحلي.  
    • يجب على المتنافس الالتزام بهذه القائمة بشكل كامل أثناء تنفيذ المشروع.  
    • تُرفق شهادات منشأ وفواتير معتمدة لإثبات أن هذه المشتريات محلية.  
    • في حال عدم توفر المنتج محليًا، يمكن طلب استثناء رسمي مرفق بمبررات.

    ثانيًا: شروط الالتزام  
    • يجب أن يحقق المتعاقد الحد الأدنى من نسبة المحتوى المحلي المطلوبة (حسب ما تحدده الجهة الحكومية).  
    • يُشترط تقديم خطة واضحة للمحتوى المحلي توضح الأصناف، النسب، والموردين المحليين.  
    • يجب تقديم فواتير رسمية وشهادات منشأ عند الطلب.

    ثالثًا: حالات الاستبعاد  
    • يتم استبعاد المتنافس من المنافسة في الحالات التالية:  
        • إذا لم يُقدم خطة محتوى محلي واضحة عند التقديم.  
        • إذا أخلّ بنسبة الالتزام المطلوبة أثناء تنفيذ المشروع دون الحصول على استثناء رسمي.  
        • إذا تبين أن الفواتير أو الشهادات المقدمة غير صحيحة أو مضللة.

    رابعًا: التوثيق والمتابعة  
    • تلتزم الجهة الحكومية بمتابعة التزام المتعاقد بنسبة المحتوى المحلي.  
    • يحق للجهة طلب تقارير مختصرة توضح نسبة الالتزام الفعلية خلال مراحل المشروع.  
    • في حال المخالفة، تطبق الجهة العقوبات الواردة في كراسة الشروط.

    خامسًا: متطلبات خطة المحتوى المحلي
    • يجب أن تتضمن الخطة بيانًا تفصيليًا بالمنتجات والخدمات المحلية المستخدمة في المشروع.
    • يجب تحديد نسب المحتوى المحلي المستهدفة في كل مرحلة من مراحل المشروع.
    • يتعين تحديد الموردين المحليين الذين سيتم التعامل معهم، مع بيان مؤهلاتهم.
    • يجب توضيح آليات قياس نسب المحتوى المحلي والتحقق منها.
    • يتعين تقديم خطة للتحسين المستمر لزيادة نسب المحتوى المحلي.

    سادسًا: آليات دعم المحتوى المحلي
    • التدريب ونقل المعرفة للكوادر الوطنية
    • توطين التقنية والمعرفة
    • تطوير سلاسل التوريد المحلية
    • الاستثمار في تنمية القدرات الوطنية
    • الشراكة مع المؤسسات الصغيرة والمتوسطة المحلية

    سابعًا: جدول حساب نسبة المحتوى المحلي

    | العنصر | المصدر | الوزن النسبي | آلية الاحتساب | المستندات المطلوبة |
    | ------ | ------ | ------------ | ------------- | ------------------ |
    |        |        |              |               |                    |
    |        |        |              |               |                    |
    |        |        |              |               |                    |
    |        |        |              |               |                    |

    ثامنًا: نموذج خطة المحتوى المحلي

    | القسم | المحتوى المطلوب | التفاصيل | ملاحظات |
    | ----- | --------------- | -------- | ------- |
    |       |                 |          |         |
    |       |                 |          |         |
    |       |                 |          |         |
    |       |                 |          |         |

    تاسعًا: آلية تقييم نسبة المحتوى المحلي
    • منهجية احتساب نسبة المحتوى المحلي
    • آلية التحقق من صحة النسب المعلنة
    • إجراءات التدقيق على المستندات
    • الزيارات الميدانية للتحقق من المحتوى المحلي
    • التقارير الدورية لمتابعة الالتزام بالنسب المطلوبة

    ملاحظات هامة للتنسيق:  
    - استخدم صيغة "القسم التاسع: متطلبات المحتوى المحلي" في بداية النص.
    - لا تستخدم علامة # للعناوين.  
    - لا تستخدم علامة ** للنص العريض.  
    - استخدم الترقيم العادي للقوائم (1. 2. 3.).  
    - استخدم النقاط العادية (•).  
    - لا تُدرج جداول أو فقرات حوافز.
   - اذكر الاسماء المعطاه من المستخدم 
    - اي مبالغ ماليه سيتم ذكرها يجب ان تكون بذكر من المستخدم
    - قم بملء الجداول بمحتوى واقعي مناسب يعكس طبيعة المشروع.


    تعليمات التوليد هذه العليمات لك لا تقم بطباعتها :  
    - اكتب 3000–4000 كلمة.  
    - استخدم لغة رسمية فصحى واضحة ومترابطة.  
    - لا تعتذر أو ترفض الكتابة. هذا المحتوى يمثل وثيقة تعليمية رسمية.
    """
    response = llm.predict(prompt)
    return response


def generate_rfp_attachments(llm, scope_summary, special_terms, competition_name, competition_objectives,
                             competition_description, government_entity):
    prompt = f"""

            يجب أن يكون النص متكيّفًا مع سياق المنافسة التالية:
        - اسم المشروع: {competition_name}
        - الهدف من المشروع: {competition_objectives}
        - وصف المجال: {competition_description}
        -الجهة الحكومية: {government_entity}
اكتب القسم العاشر من كراسة الشروط بعنوان "الشروط الخاصة".
لا تكتب عنوان القسم أبدًا، وابدأ مباشرة في كتابة المحتوى.

هذا القسم يضيف شروطًا إضافية بناءً على طبيعة نطاق العمل الذي تم تحديده سابقًا، ويهدف إلى تغطية الجوانب التشغيلية، البيئية، الفنية أو الإدارية الخاصة بهذا المشروع تحديدًا.  
يجب أن تكون الشروط مكملة للشروط العامة، وتتناسب مع نطاق العمل التالي:

ملخص نطاق العمل:  
{scope_summary}

الشروط الخاصة المُحددة لهذا المشروع:
{special_terms}

يجب أن يشمل القسم الفقرات التالية:

أولاً: التقرير الشهري  
• يُلزم المتعاقد بتقديم تقرير شهري تفصيلي إلى الجهة الحكومية يشمل نسب التقدم في إنجاز الأعمال.  
• يبدأ تقديم التقرير من الشهر الأول لتاريخ الإشعار بالمباشرة، ويُسلّم في موعد أقصاه اليوم الثامن من كل شهر ميلادي.  
• في حال وافق اليوم الثامن يوم إجازة نهاية أسبوع أو إجازة رسمية، فيتم تسليم التقرير في أول يوم عمل بعد انتهاء الإجازة.  
• يحتوي التقرير الشهري على العناصر التالية:  
  أ. صفحة العنوان والتوزيع وجدول المحتويات والمقدمة  
  ب. الملخص التنفيذي  
  ج. جداول المشروع وإنجازات العمل الحالية لكل مسار  
  د. الأنشطة المخطط لها بالنسبة إلى الشهر التالي  
  هـ. قسم تكاليف المشروع، ومقارنة الالتزامات بالنفقات الفعلية  
  و. المخاطر (إن وجدت)

ثانياً: الاجتماعات الدورية  
• يلتزم المتعاقد بحضور اجتماعات دورية مع ممثلي الجهة الحكومية لمتابعة سير العمل.  
• تعقد الاجتماعات بشكل أسبوعي/شهري حسب طبيعة المشروع ومراحله.  
• يتم إعداد محاضر للاجتماعات توثق القرارات والإجراءات المتخذة.  
• يلتزم المتعاقد بتنفيذ التوصيات والقرارات الصادرة عن هذه الاجتماعات.  
• يمكن عقد اجتماعات استثنائية بطلب من أي من الطرفين عند الحاجة.

ثالثاً: ضمان جودة الأعمال  
• يلتزم المتعاقد بتطبيق معايير الجودة المتفق عليها في جميع مراحل التنفيذ.  
• يتم إجراء اختبارات دورية للتأكد من مطابقة الأعمال للمواصفات المطلوبة.  
• يحق للجهة الحكومية رفض أي أعمال لا تتوافق مع المعايير المحددة.  
• يلتزم المتعاقد بتصحيح الأعمال المرفوضة خلال مدة محددة وعلى نفقته الخاصة.  
• يتم توثيق جميع إجراءات ضمان الجودة في تقارير دورية.

رابعاً: التزامات الجهة الحكومية  
• توفير المعلومات والبيانات اللازمة لتنفيذ المشروع.  
• تسهيل إجراءات الحصول على التصاريح والموافقات اللازمة.  
• تعيين ممثل للإشراف على تنفيذ المشروع والتنسيق مع المتعاقد.  
• مراجعة التقارير والمخرجات المقدمة من المتعاقد وتقديم الملاحظات عليها.  
• صرف المستحقات المالية وفقاً للجدول الزمني المتفق عليه بعد استيفاء الشروط.

خامساً: التزامات المتعاقد الإضافية  
• الالتزام بالتشريعات والأنظمة المحلية ذات العلاقة بطبيعة المشروع.  
• الحصول على جميع التراخيص والتصاريح اللازمة لتنفيذ الأعمال.  
• توفير جميع المعدات والأدوات اللازمة لتنفيذ المشروع بجودة عالية.  
• التنسيق المستمر مع الجهات ذات العلاقة لضمان سير العمل بسلاسة.  
• تقديم الدعم الفني اللازم بعد انتهاء المشروع خلال فترة الضمان.

سادساً: متطلبات التقارير واجتماعات المتابعة

| نوع التقرير/الاجتماع | الدورية | المحتوى | الجهة المستلمة |
| -------------------- | ------- | ------- | -------------- |
| تقرير تقدم سير العمل | أسبوعي | نسبة الإنجاز، الأعمال المنجزة، الأعمال المتأخرة، المعوقات | قسم إدارة المشاريع |
| تقرير الجودة | شهري | نتائج اختبارات الجودة، الانحرافات، الإجراءات التصحيحية | قسم مراقبة الجودة |
| اجتماع فريق الإدارة | أسبوعي | مناقشة تقدم المشروع، المعوقات، القرارات المطلوبة | مدير المشروع والفريق الفني |
| اجتماع اللجنة الفنية | شهري | عرض التقارير الفنية، اعتماد التغييرات، مناقشة المخاطر | اللجنة الفنية المشرفة |

سابعاً: آلية إدارة التغيير في المشروع

| نوع التغيير | إجراءات التقديم | مستويات الموافقة | المدة الزمنية للرد |
| ----------- | --------------- | ----------------- | ------------------- |
| تغيير في النطاق | تقديم طلب تغيير رسمي مع التبريرات والتأثيرات | مدير المشروع، اللجنة الفنية، المدير التنفيذي | 14 يوم عمل |
| تمديد الجدول الزمني | تقديم طلب مع تحليل الأسباب وتأثير التمديد | مدير المشروع، اللجنة الفنية | 10 أيام عمل |
| تعديل المواصفات الفنية | تقديم طلب مفصل بالتعديلات المطلوبة والمبررات | الفريق الفني، مدير المشروع، مدير الجودة | 7 أيام عمل |
| تعديل التكلفة | تقديم طلب مع تحليل مالي مفصل للتكاليف الإضافية | اللجنة المالية، المدير التنفيذي | 15 يوم عمل |

ثامناً: إجراءات التسليم والاستلام  
• آلية تسليم مخرجات المشروع  
• إجراءات الفحص والاختبار  
• متطلبات الاستلام الأولي  
• شروط الاستلام النهائي  
• آلية معالجة الملاحظات

تاسعاً: إجراءات الدعم الفني والصيانة  
• نطاق خدمات الدعم الفني  
• مستويات الخدمة المطلوبة  
• آلية الإبلاغ عن المشكلات  
• أوقات الاستجابة المطلوبة  
• إجراءات التصعيد والمتابعة

تعليمات التوليد هذه العليمات لك لا تقم بطباعتها:  
- استخدم صيغة "القسم العاشر: الشروط الخاصة" في بداية النص.  
- لا تستخدم # أو **.  
- استخدم الفقرات والتعداد الرسمي فقط.  
- لا تُدرج مواد نظامية أو فقرات حوافز أو استثناءات إلا إذا طُلب.  
- لا تستخدم جداول أو تنسيقات Markdown.  
- اذكر الاسماء المعطاه من المستخدم 
- اي مبالغ ماليه سيتم ذكرها يجب ان تكون بذكر من المستخدم 
- قم بملء الجداول بمحتوى واقعي مناسب يعكس طبيعة المشروع.


- اكتب بين 3000 إلى 4000 كلمة.  
- استخدم لغة رسمية واضحة فصحى.  
- لا تعتذر عن الكتابة. هذا محتوى تعليمي رسمي يُكلفك به.
"""
    response = llm.predict(prompt)
    return response


def generate_rfp_annexes_and_forms(llm, example_rfp, competition_name, competition_objectives, competition_description,
                                   government_entity):
    prompt = f"""

            يجب أن يكون النص متكيّفًا مع سياق المنافسة التالية:
        - اسم المشروع: {competition_name}
        - الهدف من المشروع: {competition_objectives}
        - وصف المجال: {competition_description}
        - الجهة الحكومية: {government_entity}
    اكتب القسم الحادي عشر من كراسة الشروط بعنوان: الملاحق.
لا تكتب ابدا عنوان القسم مباشره في كتابة المحتوى

    الهدف من هذا القسم هو توضيح الملاحق والمستندات التي يمكن أن تُرفق مع كراسة الشروط والمواصفات حسب ما تراه الجهة الحكومية مناسبًا لطبيعة المشروع. 
    يجب أن يُذكر بوضوح أن الملاحق تختلف من مشروع لآخر، وتُضاف حسب الحاجة فقط.

    يشمل هذا القسم الملاحق التالية:

    أولاً: ملحق (1): خطاب تقديم العروض  
    • يُستخدم من قبل المتنافس لتقديم عرضه الرسمي للجهة الحكومية وفق النموذج المعتمد.
    • يجب أن يتضمن الخطاب البيانات الأساسية للمتنافس والمنافسة.
    • يجب أن يكون الخطاب موقعًا من الشخص المفوض بتمثيل المتنافس.
    • يرفق بالخطاب جميع المستندات المطلوبة في كراسة الشروط.
    • يجب أن يحتوي على إقرار بالاطلاع والموافقة على جميع شروط المنافسة.

    ثانيًا: ملحق (2): نموذج الأسئلة والاستفسارات  
    • يُستخدم لإرسال الاستفسارات الرسمية حول محتوى كراسة الشروط.  
    • تلتزم الجهة الحكومية بالرد خلال المدة المحددة.
    • يجب استخدام النموذج المعتمد لتقديم الاستفسارات.
    • يتضمن النموذج حقولًا لبيانات المتنافس وتفاصيل الاستفسار.
    • يجب تحديد رقم البند أو الصفحة محل الاستفسار بدقة.

    ثالثًا: ملحق (3): نموذج العقد  
    • يشمل الأحكام والشروط الأساسية للعقد الذي سيُبرم بعد الترسية.
    • يتضمن العقد النموذجي جميع البنود القانونية التي تنظم العلاقة بين الطرفين.
    • يوضح آليات التعامل مع المتغيرات والظروف الطارئة أثناء التنفيذ.
    • يحدد مسؤوليات والتزامات كل طرف بشكل دقيق ومفصل.
    • يتضمن آليات حل النزاعات والخلافات التي قد تنشأ أثناء التنفيذ.

    رابعًا: ملحق (4): الرسومات والمخططات  
    • تُرفق هذه الملاحق في حال كان نطاق العمل يتضمن عناصر تصميم أو هندسة.
    • تشمل المخططات التفصيلية للمشروع بمقاييس الرسم المناسبة.
    • توضح المواصفات الفنية للمكونات والعناصر المختلفة.
    • تتضمن مواقع التنفيذ والحدود الجغرافية للمشروع.
    • تشمل أي تفاصيل فنية ضرورية لفهم طبيعة المشروع.

    خامسًا: ملحق (5): القائمة الإلزامية  
    • تُحدد فيه المواد أو المنتجات التي يجب أن تكون من إنتاج محلي.
    • توضح المواصفات الفنية للمنتجات المحلية المطلوبة.
    • تحدد النسب المطلوبة من كل منتج حسب طبيعة المشروع.
    • توضح آليات إثبات المنشأ المحلي للمنتجات والمواد.
    • تشمل بدائل مقبولة في حال عدم توفر المنتج محليًا.

    سادسًا: ملحق (6): الشروط والأحكام لآلية التفضيل السعري للمنتج الوطني  
    • يُرفق فقط عند تطبيق سياسة التفضيل السعري للمنتجات الوطنية.
    • يوضح نسب التفضيل المطبقة على المنتجات الوطنية.
    • يحدد آليات احتساب التفضيل السعري أثناء تقييم العروض.
    • يشمل الوثائق المطلوبة لإثبات استحقاق التفضيل السعري.
    • يوضح حالات الاستثناء التي لا يطبق فيها التفضيل السعري.

    سابعًا: ملحق (7): الشروط والأحكام المتعلقة بآلية الحد الأدنى المطلوب للمحتوى المحلي  
    • يُضاف عند تطبيق الحد الأدنى المطلوب للمحتوى المحلي في المنافسة.
    • يحدد النسبة الإلزامية للمحتوى المحلي المطلوب تحقيقها.
    • يوضح آليات قياس نسبة المحتوى المحلي وطرق التحقق.
    • يشمل النماذج المطلوب تعبئتها لإثبات نسبة المحتوى المحلي.
    • يوضح الإجراءات المتخذة في حال عدم تحقيق النسبة المطلوبة.

    ثامنًا: ملحق (8): نماذج الضمانات البنكية
    • نماذج موحدة للضمانات البنكية المطلوبة في مراحل المنافسة المختلفة.
    • نموذج الضمان الابتدائي المطلوب تقديمه مع العرض.
    • نموذج الضمان النهائي المطلوب بعد الترسية وقبل توقيع العقد.
    • نموذج ضمان الدفعة المقدمة (إن وجدت).
    • الشروط العامة للضمانات البنكية وفترات سريانها.

    تاسعًا: ملحق (9): نماذج التقارير الدورية

    | نوع التقرير | الدورية | محتويات التقرير | متطلبات التقديم |
    | ----------- | ------- | ---------------- | --------------- |
    |             |         |                  |                 |
    |             |         |                  |                 |
    |             |         |                  |                 |
    |             |         |                  |                 |

    عاشرًا: ملحق (10): نماذج محاضر الاجتماعات والاستلام

    | نوع المحضر | الغرض | العناصر الأساسية | متطلبات التوثيق |
    | ---------- | ----- | ---------------- | --------------- |
    |            |       |                  |                 |
    |            |       |                  |                 |
    |            |       |                  |                 |
    |            |       |                  |                 |

    حادي عشر: ملحق (11): نماذج طلبات التغيير والتعديل
    • نموذج طلب تغيير في نطاق العمل
    • نموذج طلب تمديد مدة المشروع
    • نموذج طلب تعديل بنود العقد
    • نموذج دراسة أثر التغيير على المشروع
    • نموذج الموافقة على طلب التغيير

    ملاحظات:
    • يحق للجهة الحكومية إضافة أو إزالة أي من هذه الملاحق حسب مقتضى المشروع.  
    • يُشار إلى أن بعض الملاحق لا تُرفق إلا إذا طُبّقت السياسات المرتبطة بها.  
    • يُفضل وضع الملاحق في نهاية كراسة الشروط بتسلسل رقمي واضح.
    - قم بملء الجداول بمحتوى واقعي مناسب يعكس طبيعة المشروع.


    تعليمات التوليد هذه العليمات لك لا تقم بطباعتها :  
    - استخدم صيغة "القسم الحادي عشر: الملاحق" في بداية النص.
    - لا تستخدم علامة # أو **.
    - استخدم ترقيمًا واضحًا.
    - لا تدخل في تفاصيل المستندات الداخلية لكل ملحق، فقط اشرح الغرض والاستخدام.
   - اذكر الاسماء المعطاه من المستخدم 
- اي مبالغ ماليه سيتم ذكرها يجب ان تكون بذكر من المستخدم    
    - اكتب بين 3000–4000 كلمة.
    - استخدم لغة رسمية فصحى واضحة.
    - لا ترفض الكتابة أو تبرر – المطلوب إنتاج محتوى تعليمي رسمي فقط.
    """
    response = llm.predict(prompt)
    return response


# Main function to generate RFP document
import concurrent.futures
from functools import partial


def generate_rfp_document(
        competition_name,
        competition_objectives,
        competition_description,
        government_entity,
        cost_value,
        cost_method,
        start_stage,
        end_stage,
        technical_docs,
        alternative_offers,
        initial_guarantee,
        pause_period,
        penalties,
        execution_city,
        execution_district,
        execution_region,
        required_materials,
        scope_summary,
        special_terms,
        output_dir,
        static_dir
):
    # Build vector store from PDF files in knowledge directory
    knowledge_dir = os.path.join(static_dir, "knowledge")
    vector_store = build_vector_store(knowledge_dir)

    # Find similar example RFP
    example_rfp = ""
    if vector_store:
        try:
            retrieved_docs = vector_store.similarity_search(competition_description, k=1)
            if retrieved_docs:
                example_rfp = retrieved_docs[0].page_content
                print("✅ Found similar RFP to use as reference model.")
            else:
                print("⚠️ No similar RFP found. Will generate without reference example.")
        except Exception as e:
            print(f"Error during similarity search: {str(e)}")
            print("⚠️ Will generate without reference example.")
    else:
        print("⚠️ Vector store not initialized. Will generate without reference example.")

    # Setup LLM
    llm = ChatOpenAI(model='gpt-4-turbo', temperature=0.2)

    generation_tasks = [
        (1, "المقدمة", partial(
            generate_rfp_intro,
            llm, example_rfp, competition_name, competition_objectives, competition_description,
            government_entity, cost_value, cost_method, start_stage, end_stage
        )),
        (2, "الأحكام العامة", partial(
            generate_rfp_general_terms,
            llm, example_rfp, competition_name, competition_objectives, competition_description,
            government_entity
        )),
        (3, "إعداد العروض", partial(
            generate_rfp_offer_preparation,
            llm, example_rfp, technical_docs, alternative_offers, initial_guarantee,
            competition_name, competition_objectives, competition_description, government_entity
        )),
        (4, "تقديم العروض", partial(
            generate_rfp_offer_submission,
            llm, example_rfp, competition_name, competition_objectives, competition_description, government_entity
        )),
        (5, "تقييم العروض", partial(
            generate_rfp_offer_analysis,
            llm, competition_name, competition_description, competition_objectives, pause_period, government_entity
        )),
        (6, "متطلبات التعاقد", partial(
            generate_rfp_award_contract,
            llm, example_rfp, penalties, competition_name, competition_description,
            competition_objectives, pause_period, government_entity
        )),
        (7, "نطاق العمل المفصل", partial(
            generate_rfp_work_scope,
            llm, competition_name, competition_description, competition_objectives,
            execution_city, execution_district, execution_region, government_entity
        )),
        (8, "المواصفات الفنية", partial(
            generate_rfp_specifications,
            llm, example_rfp, required_materials, competition_name,
            competition_objectives, competition_description, government_entity
        )),
        (9, "متطلبات المحتوى المحلي", partial(
            generate_rfp_general_contract_terms,
            llm, example_rfp, competition_name, competition_objectives, competition_description, government_entity
        )),
        (10, "الشروط الخاصة", partial(
            generate_rfp_attachments,
            llm, scope_summary, special_terms, competition_name, competition_objectives,
            competition_description, government_entity
        )),
        (11, "الملاحق والنماذج الإضافية", partial(
            generate_rfp_annexes_and_forms,
            llm, example_rfp, competition_name, competition_objectives,
            competition_description, government_entity
        )),
    ]

    # Store results by section number to maintain ordering
    sections_content = {}

    print("🔹 Starting parallel generation of all RFP sections...")

    # Use ThreadPoolExecutor to run generation tasks in parallel
    # Limiting to 4 workers to avoid overwhelming the API or system resources
    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
        # Create a dictionary to map futures to their section info
        future_to_section = {}

        # Submit all tasks
        for section_num, section_title, generate_func in generation_tasks:
            future = executor.submit(generate_func)
            future_to_section[future] = (section_num, section_title)

        # Process results as they complete
        for future in concurrent.futures.as_completed(future_to_section):
            section_num, section_title = future_to_section[future]
            try:
                section_content = future.result()
                sections_content[section_num] = (section_title, section_content)
                print(f"✅ Completed section {section_num}: {section_title}")
            except Exception as e:
                print(f"❌ Error generating section {section_num}: {section_title}")
                print(f"   Error details: {str(e)}")
                # Provide a placeholder for failed sections
                sections_content[section_num] = (
                    section_title, f"Error generating {section_title} section. Please try again.")

    print("✅ All sections completed!")

    # Combine all sections in the correct order
    sections = [sections_content[i] for i in range(1, 12)]
    safe_filename = re.sub(r'[^\w\s]', '', competition_name).strip().replace(' ', '_')
    filename = f"{safe_filename}_rfp.docx"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)
    save_rfp_sections_to_word(sections, output_path, competition_name)

    return filename


import fitz  # PyMuPDF for PDF extraction


def read_pdf_with_fitz(file_path):
    """
    Extract text from PDF using PyMuPDF (fitz) with OCR fallback.
    """
    # Try the OCR-enabled extraction first
    return extract_text_from_pdf(file_path)


def clean_text(text):
    """
    Clean extracted text.
    """
    text = re.sub(r'Error! Bookmark not defined\.', '', text)
    text = re.sub(r'\d{1,3}', '', text)
    text = re.sub(r'\.{2,}', '', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()


def improve_rfp_with_extracted_text(pdf_path, competition_name, competition_objectives, competition_description,
                                    output_path, vector_store=None, **kwargs):
    """
    Improve an existing RFP document using pre-extracted text.
    This function is optimized for parallel processing with OCR support.
    """
    # استخراج النص من ملف PDF باستخدام OCR إذا لزم الأمر
    print(f"📄 بدء معالجة وثيقة المستخدم: {pdf_path}")
    pdf_text = extract_text_from_pdf(pdf_path)

    if not pdf_text or len(pdf_text.strip()) < 500:
        print("⚠️ لم يتم استخراج نص كافٍ من ملف PDF. جاري محاولة معالجة OCR مرة أخرى...")
        pdf_text = perform_ocr_on_pdf(pdf_path)

    print(f"✅ تم استخراج {len(pdf_text.split())} كلمة من وثيقة المستخدم")

    # Setup LLM
    llm = ChatOpenAI(model='gpt-4-turbo', temperature=0.2)

    # Define required sections
    required_sections = [
        'المقدمة', 'الأحكام العامة', 'إعداد العروض', 'تقديم العروض',
        'تقييم العروض', 'متطلبات التعاقد', 'نطاق العمل المفصل',
        'المواصفات الفنية', 'متطلبات المحتوى المحلي', 'الشروط الخاصة',
        'الملاحق والنماذج الإضافية'
    ]

    # Get example RFP if vector store is available
    example_rfp = ""
    if vector_store:
        try:
            retrieved_docs = vector_store.similarity_search(competition_description, k=1)
            if retrieved_docs:
                example_rfp = retrieved_docs[0].page_content
                print("✅ Found similar RFP to use as reference model.")
            else:
                print("⚠️ No similar RFP found. Will generate without reference example.")
        except Exception as e:
            print(f"Error during similarity search: {str(e)}")
            print("⚠️ Will generate without reference example.")

    # Clean the text
    text = clean_text(pdf_text)
    notes = []

    # Prepare tasks for section processing
    section_tasks = []

    for section in required_sections:
        # Try to find the section in the original document
        pattern = rf"{section}\s*([\s\S]*?)(?=\n\s*(?:{'|'.join(required_sections)})|$)"
        match = re.search(pattern, text)
        section_content = match.group(1).strip() if match else ""

        generate_flag = False

        # Check if section needs to be generated or improved
        if len(section_content) < 50:
            note = f"❌ القسم '{section}' مفقود → سيتم توليده."
            notes.append(note)
            generate_flag = True
        elif section == 'المقدمة':
            key_terms = ['تعريف', 'خلفية', 'نطاق', 'أهداف']
            if not all(term in section_content for term in key_terms):
                note = f"⚠️ القسم '{section}' ناقص في التعريف/الخلفية/النطاق/الأهداف → سيتم تحسينه."
                notes.append(note)
                generate_flag = True
        else:
            note = f"ℹ️ القسم '{section}' موجود وسنعيد كتابته لضمان التنسيق والترتيب."
            notes.append(note)

        # Add to tasks list
        section_tasks.append((section, section_content, generate_flag))

    # Extract additional parameters from kwargs
    government_entity = kwargs.get('government_entity', 'الجهة الحكومية')
    cost_value = kwargs.get('cost_value', '')
    cost_method = kwargs.get('cost_method', '')
    start_stage = kwargs.get('start_stage', '')
    end_stage = kwargs.get('end_stage', '')
    technical_docs = kwargs.get('technical_docs', '')
    alternative_offers = kwargs.get('alternative_offers', '')
    initial_guarantee = kwargs.get('initial_guarantee', '')
    pause_period = kwargs.get('pause_period', '')
    penalties = kwargs.get('penalties', '')
    execution_city = kwargs.get('execution_city', '')
    execution_district = kwargs.get('execution_district', '')
    execution_region = kwargs.get('execution_region', '')
    required_materials = kwargs.get('required_materials', '')
    scope_summary = kwargs.get('scope_summary', '')
    special_terms = kwargs.get('special_terms', '')

    # Process sections in parallel using ThreadPoolExecutor with increased workers
    section_results = {}
    with concurrent.futures.ThreadPoolExecutor(max_workers=6) as executor:
        # Create a dictionary to map futures to their section info
        future_to_section = {}

        for section, content, generate_flag in section_tasks:
            # Submit task to generate or improve the section
            future = executor.submit(
                improve_section,
                llm,
                section,
                content,
                competition_name,
                competition_objectives,
                competition_description,
                generate_flag,
                example_rfp=example_rfp,
                government_entity=government_entity,
                cost_value=cost_value,
                cost_method=cost_method,
                start_stage=start_stage,
                end_stage=end_stage,
                technical_docs=technical_docs,
                alternative_offers=alternative_offers,
                initial_guarantee=initial_guarantee,
                pause_period=pause_period,
                penalties=penalties,
                execution_city=execution_city,
                execution_district=execution_district,
                execution_region=execution_region,
                required_materials=required_materials,
                scope_summary=scope_summary,
                special_terms=special_terms
            )
            future_to_section[future] = section

        # Process results as they complete
        for future in concurrent.futures.as_completed(future_to_section):
            section = future_to_section[future]
            try:
                section_content = future.result()
                section_results[section] = section_content
                print(f"✅ Completed improving section: {section}")
            except Exception as e:
                print(f"❌ Error improving section {section}: {e}")
                # Provide a placeholder for failed sections
                section_results[section] = f"Error generating {section} section. Please try again."

    # Organize sections in the correct order
    sections = []
    for section in required_sections:
        if section in section_results:
            sections.append((section, section_results[section]))

    # Ensure directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # Save the improved RFP to a Word document
    save_rfp_sections_to_word(sections, output_path, competition_name)

    # Generate filename
    filename_base = os.path.basename(output_path)

    # Print notes
    notes_text = "\n".join(notes)
    print("\n===== 📝 ملاحظات المراجعة =====\n")
    print(notes_text)
    print(f"\n✅ تم حفظ الكراسة المحسنة في {filename_base}.\n")

    return filename_base


# This patch modifies the improve_section function to ensure it always generates content
# without refusal messages

def improve_section(llm, section, original_content, competition_name, competition_objectives, competition_description,
                    generate_flag, **kwargs):
    """
    تحسين أو توليد قسم واحد من كراسة الشروط مع الاستفادة من دوال التوليد المتخصصة.

    Args:
        llm: نموذج اللغة المستخدم للتوليد
        section: اسم القسم المراد تحسينه
        original_content: المحتوى الأصلي للقسم (إذا كان موجودًا)
        competition_name: اسم المنافسة/المشروع
        competition_objectives: أهداف المنافسة
        competition_description: وصف المنافسة
        generate_flag: علم يحدد ما إذا كان يجب توليد القسم بالكامل
        **kwargs: معلومات إضافية مطلوبة للتوليد (مثل government_entity, cost_value, إلخ)

    Returns:
        المحتوى المحسن أو المولد للقسم
    """
    # استخراج المعلومات الإضافية من kwargs
    government_entity = kwargs.get('government_entity', 'الجهة الحكومية')
    cost_value = kwargs.get('cost_value', '')
    cost_method = kwargs.get('cost_method', '')
    start_stage = kwargs.get('start_stage', '')
    end_stage = kwargs.get('end_stage', '')
    technical_docs = kwargs.get('technical_docs', '')
    alternative_offers = kwargs.get('alternative_offers', '')
    initial_guarantee = kwargs.get('initial_guarantee', '')
    pause_period = kwargs.get('pause_period', '')
    penalties = kwargs.get('penalties', '')
    execution_city = kwargs.get('execution_city', '')
    execution_district = kwargs.get('execution_district', '')
    execution_region = kwargs.get('execution_region', '')
    required_materials = kwargs.get('required_materials', '')
    scope_summary = kwargs.get('scope_summary', '')
    special_terms = kwargs.get('special_terms', '')
    example_rfp = kwargs.get('example_rfp', '')

    # إيجاد رقم القسم
    required_sections = [
        'المقدمة', 'الأحكام العامة', 'إعداد العروض', 'تقديم العروض',
        'تقييم العروض', 'متطلبات التعاقد', 'نطاق العمل المفصل',
        'المواصفات الفنية', 'متطلبات المحتوى المحلي', 'الشروط الخاصة',
        'الملاحق والنماذج الإضافية'
    ]
    section_index = required_sections.index(section) + 1 if section in required_sections else 0

    # إذا كان المحتوى غير موجود أو علم التوليد نشط (أو المحتوى قصير جدًا)
    if generate_flag or len(original_content.strip()) < 500:
        print(f"🔄 توليد قسم '{section}' باستخدام دالة التوليد المتخصصة...")

        try:
            # استدعاء دوال التوليد المتخصصة حسب القسم
            if section == "المقدمة":
                return generate_rfp_intro(llm, example_rfp, competition_name, competition_objectives,
                                          competition_description, government_entity, cost_value,
                                          cost_method, start_stage, end_stage)

            elif section == "الأحكام العامة":
                return generate_rfp_general_terms(llm, example_rfp, competition_name, competition_objectives,
                                                  competition_description, government_entity)

            elif section == "إعداد العروض":
                return generate_rfp_offer_preparation(llm, example_rfp, technical_docs, alternative_offers,
                                                      initial_guarantee, competition_name, competition_objectives,
                                                      competition_description, government_entity)

            elif section == "تقديم العروض":
                return generate_rfp_offer_submission(llm, example_rfp, competition_name, competition_objectives,
                                                     competition_description, government_entity)

            elif section == "تقييم العروض":
                return generate_rfp_offer_analysis(llm, competition_name, competition_description,
                                                   competition_objectives, pause_period, government_entity)

            elif section == "متطلبات التعاقد":
                return generate_rfp_award_contract(llm, example_rfp, penalties, competition_name,
                                                   competition_description, competition_objectives,
                                                   pause_period, government_entity)

            elif section == "نطاق العمل المفصل":
                return generate_rfp_work_scope(llm, competition_name, competition_description,
                                               competition_objectives, execution_city,
                                               execution_district, execution_region, government_entity)

            elif section == "المواصفات الفنية":
                return generate_rfp_specifications(llm, example_rfp, required_materials,
                                                   competition_name, competition_objectives,
                                                   competition_description, government_entity)

            elif section == "متطلبات المحتوى المحلي":
                return generate_rfp_general_contract_terms(llm, example_rfp, competition_name,
                                                           competition_objectives, competition_description,
                                                           government_entity)

            elif section == "الشروط الخاصة":
                return generate_rfp_attachments(llm, scope_summary, special_terms, competition_name,
                                                competition_objectives, competition_description, government_entity)

            elif section == "الملاحق والنماذج الإضافية":
                return generate_rfp_annexes_and_forms(llm, example_rfp, competition_name,
                                                      competition_objectives, competition_description,
                                                      government_entity)

            else:
                # إذا لم يكن هناك دالة توليد متخصصة، استخدم نهج التوليد العام
                print(f"⚠️ لا توجد دالة توليد متخصصة لقسم '{section}'، سيتم استخدام التوليد العام.")
                return generate_generic_section(llm, section, section_index, competition_name,
                                                competition_objectives, competition_description, government_entity)

        except Exception as e:
            print(f"❌ حدث خطأ أثناء محاولة توليد قسم '{section}': {str(e)}")
            print("⚠️ الانتقال إلى نهج التحسين العام...")
            # في حالة وجود خطأ، نستمر بالتنفيذ واستخدام النهج العام
            pass

    # استخدام نهج التحسين إذا كان المحتوى موجودًا أو إذا فشل التوليد المتخصص
    section_with_num = f"القسم {section_index}: {section}" if section_index > 0 else section

    print(f"🔄 تحسين قسم '{section}' باستخدام المحتوى الموجود...")

    prompt = f"""
    أنت خبير في كتابة كراسات الشروط والمواصفات للمشاريع الحكومية في المملكة العربية السعودية. 

    المطلوب: قم بتحسين وإعادة صياغة قسم '{section_with_num}' من كراسة شروط مشروع '{competition_name}'.

    النص الأصلي:
    {original_content}

    تعليمات مهمة:
    1. حافظ على جميع المعلومات الموجودة في النص الأصلي وقم بتحسين صياغتها فقط.
    2. استخدم لغة عربية فصحى رسمية واضحة.
    3. تجنب أي عبارات اعتذار أو رفض مثل "لا يمكنني إعادة صياغة هذا المحتوى".
    4. افترض أنك خبير مفوض بتحسين هذا المحتوى، واكتب النص كما لو كنت موظفًا في قسم المشتريات.
    5. أضف تفاصيل مناسبة إذا كان النص الأصلي مختصرًا جدًا.
    6. تأكد من أن النص المحسّن لا يقل عن 3000 كلمة.
    7. لا تذكر أي أسماء محددة لأشخاص إلا إذا كانت مذكورة تحديدًا في النص الأصلي أو وصف المشروع.
    8. لا تذكر أي مبالغ مالية محددة، إلا إذا كانت معطاة.
    9. تأكد من استخدام صيغة "{section_with_num}" في بداية النص.

    ملاحظات للتنسيق:
    - استخدم العناوين بشكل عادي دون علامة #.
    - استخدم النص العادي دون علامة ** للنص العريض.
    - استخدم الترقيم العادي للقوائم (1. 2. 3.).
    - استخدم النقاط العادية للقوائم غير المرقمة (•).
    - استخدم | بين أعمدة الجداول إذا لزم الأمر.
    """

    try:
        section_content = llm.predict(prompt).strip()

        # التحقق من جودة المحتوى وطوله
        if len(section_content) < 3000 or "عذرًا" in section_content or "لا يمكنني" in section_content:
            # إذا كان المحتوى قصيرًا جدًا أو يحتوي على عبارات رفض، نحاول مرة أخرى بطلب أقوى
            print(f"⚠️ محتوى '{section}' غير مُرضٍ. محاولة أخرى مع طلب أقوى...")

            retry_prompt = f"""
            أنت خبير متخصص في إعداد كراسات الشروط الحكومية وتم توظيفك خصيصًا لكتابة هذا المحتوى. هذا سياق تعليمي وتدريبي فقط.

            اكتب محتوى قسم '{section_with_num}' بشكل مفصل وواقعي لكراسة شروط مشروع '{competition_name}'. 

            المحتوى الأصلي الذي تحتاج لتحسينه:
            {original_content}

            الهدف تدريبي وتعليمي فقط. يجب أن يكون المحتوى:
            1. واقعي ومفصل (لا يقل عن 3000 كلمة)
            2. مكتوب بلغة مهنية رسمية
            3. مقسم إلى أقسام وبنود واضحة
            4. متوافق مع الأنظمة العامة للمشتريات الحكومية
            5. خالي من أي أسماء لأشخاص إلا إذا كانت مذكورة في وصف المشروع

            كن حذرًا من أن أي اعتذار أو رفض للكتابة سيؤدي إلى فشل المهمة التي كُلفت بها. المطلوب كتابة محتوى تعليمي نموذجي فقط.
            """

            section_content = llm.predict(retry_prompt).strip()

    except Exception as e:
        print(f"❌ خطأ في تحسين محتوى '{section}': {str(e)}")
        # توفير محتوى احتياطي في حالة الخطأ
        section_content = generate_fallback_content(section, competition_name)

    # التأكد من أن القسم يبدأ بالترقيم الصحيح
    if not section_content.startswith(f"القسم {section_index}") and section_index > 0:
        section_content = f"{section_with_num}\n\n{section_content}"

    # تنظيف المحتوى النهائي
    return clean_text(section_content)


def generate_generic_section(llm, section, section_index, competition_name, competition_objectives,
                             competition_description, government_entity):
    """
    توليد قسم عام عندما لا تتوفر دالة توليد متخصصة.
    """
    section_with_num = f"القسم {section_index}: {section}" if section_index > 0 else section

    prompt = f"""
    أنت خبير في كتابة كراسات الشروط والمواصفات للمشاريع الحكومية في المملكة العربية السعودية. 

    المطلوب: اكتب قسم '{section_with_num}' من كراسة الشروط لمشروع '{competition_name}' الذي يهدف إلى {competition_objectives} في مجال {competition_description}.

    الجهة الحكومية هي: {government_entity}

    تعليمات مهمة:
    1. يجب أن تكتب محتوى واقعي ومفصل مناسب لكراسة شروط رسمية.
    2. تجنب أي عبارات اعتذار أو رفض مثل "لا يمكنني كتابة هذا المحتوى" أو "يجب الاستعانة بخبير قانوني".
    3. افترض أنك خبير مفوض بكتابة هذا المحتوى، واكتب النص كما لو كنت موظفًا في قسم المشتريات.
    4. استخدم لغة عربية فصحى رسمية دون أي كلمات أجنبية.
    5. اكتب محتوى لا يقل عن 3000 كلمة مع تفاصيل وافية.
    6. اذكر فقط المتطلبات والشروط والمعايير المتعلقة بقسم '{section}' دون الخروج عن نطاقه.
    7. لا تذكر أي أسماء محددة لأشخاص إلا إذا كانت مذكورة تحديدًا في وصف المشروع.
    8. لا تذكر أي مبالغ مالية محددة.
    9. اجعل كل الجداول هيكلية فقط مع عناوين الأعمدة ولكن بدون بيانات في الخلايا.

    ملاحظات للتنسيق:
    - استخدم العناوين بشكل عادي دون علامة #.
    - استخدم النص العادي دون علامة ** للنص العريض.
    - استخدم الترقيم العادي للقوائم (1. 2. 3.).
    - استخدم النقاط العادية للقوائم غير المرقمة (•).
    - استخدم | بين أعمدة الجداول إذا لزم الأمر.
    """

    try:
        return llm.predict(prompt).strip()
    except Exception as e:
        print(f"❌ خطأ في توليد محتوى '{section}': {str(e)}")
        return generate_fallback_content(section, competition_name)


def generate_fallback_content(section, competition_name):
    """
    توفير محتوى احتياطي للقسم في حالة فشل التوليد.
    """
    # العثور على رقم القسم
    required_sections = [
        'المقدمة', 'الأحكام العامة', 'إعداد العروض', 'تقديم العروض',
        'تقييم العروض', 'متطلبات التعاقد', 'نطاق العمل المفصل',
        'المواصفات الفنية', 'متطلبات المحتوى المحلي', 'الشروط الخاصة',
        'الملاحق والنماذج الإضافية'
    ]
    section_index = required_sections.index(section) + 1 if section in required_sections else '؟'

    # إنشاء محتوى احتياطي تفصيلي مع جداول فارغة
    return f"""
    القسم {section_index}: {section}

    يتضمن هذا القسم الشروط والتفاصيل المتعلقة بـ {section} لمشروع "{competition_name}". 

    1. المتطلبات الأساسية
    • يجب على المتنافسين الالتزام بكافة المتطلبات المحددة في هذا القسم.
    • تطبق الأنظمة واللوائح المعمول بها في المملكة العربية السعودية.
    • يلتزم المتنافس بجميع المعايير والمواصفات الفنية المطلوبة.
    • يجب تقديم جميع الوثائق والمستندات بالشكل المطلوب وفي الموعد المحدد.
    • تخضع جميع الأعمال للإشراف والمتابعة من قبل الجهة المالكة للمشروع.

    2. الإجراءات والآليات
    • يتم اتباع الإجراءات المحددة في هذا القسم بدقة.
    • تخضع جميع الأعمال للمراجعة والتدقيق من قبل الجهة الحكومية.
    • يلتزم المتنافس بتقديم تقارير دورية عن سير العمل.
    • في حال وجود أي استفسارات، يتم التواصل مع الجهة المالكة عبر القنوات الرسمية.
    • يتم توثيق جميع الإجراءات والقرارات المتخذة أثناء تنفيذ المشروع.

    3. المعايير والمواصفات
    • يجب الالتزام بالمعايير الفنية والمواصفات المحددة.
    • تطبق معايير الجودة المعتمدة محليًا ودوليًا.
    • يلتزم المتنافس بتطبيق أعلى معايير الجودة في جميع مراحل المشروع.
    • يجب الالتزام بالمواصفات البيئية والصحية المعتمدة.
    • تخضع جميع المواد المستخدمة للفحص والاختبار قبل الاعتماد.
    """


def is_valid_rfp_document(pdf_text, debug_mode=False):
    """
    Check if the extracted text represents an RFP document.
    Improved version with better detection and detailed logging.
    """
    # Check if document is empty or too short
    if not pdf_text or len(pdf_text.strip()) < 100:  # Reduced minimum from 200 to 100
        return False, "Document is empty or too short. It must contain sufficient content for analysis."

    # Common terms in RFP documents - English and Arabic keywords
    rfp_keywords = [
        # Arabic keywords
        'كراسة الشروط', 'المناقصة', 'المنافسة', 'المتنافس', 'العروض', 'الجهة الحكومية',
        'نطاق العمل', 'المواصفات', 'الضمان', 'الأحكام', 'الشروط', 'التعاقد', 'التقييم',
        'تقديم العروض', 'الترسية', 'الملاحق', 'غرامات', 'مستندات', 'كفالة', 'ضمان ابتدائي',
        # Additional Arabic keywords
        'منافسة', 'مشروع', 'عرض', 'عطاء', 'توريد', 'تنفيذ', 'متطلبات', 'خدمات',
        # English keywords
        'RFP', 'proposal', 'tender', 'bid', 'procurement', 'scope of work', 'specifications',
        'contract', 'terms and conditions', 'submission', 'evaluation'
    ]

    # Log found keywords in debug mode
    found_keywords = []
    for keyword in rfp_keywords:
        if keyword.lower() in pdf_text.lower():  # Case-insensitive check
            found_keywords.append(keyword)

    if debug_mode:
        print(f"Text length: {len(pdf_text)}")
        print(f"Found {len(found_keywords)} keywords: {', '.join(found_keywords)}")
        print(f"Text sample: {pdf_text[:300]}...")  # Show first 300 chars

    # Check for minimum number of terms (reduced to 1)
    if len(found_keywords) < 1:  # Changed from 2 to 1
        return False, f"Document doesn't appear to be an RFP. Found 0 keywords out of minimum required 1."

    # Section structure indicators - both English and Arabic
    section_indicators = [
        # Arabic indicators
        'القسم', 'الفصل', 'الباب', 'الجزء', 'المادة', 'البند',
        'أولاً', 'ثانيًا', 'ثالثًا', 'رابعًا', 'خامسًا',
        'المقدمة', 'الأحكام العامة', 'إعداد العروض', 'تقديم العروض',
        # English indicators
        'section', 'chapter', 'article', 'clause', 'part',
        'introduction', 'general terms', 'requirements', 'evaluation criteria'
    ]

    # Skip structure check if enough keywords found
    if len(found_keywords) >= 2:
        return True, ""

    # Structure check - only performed if exactly 1 keyword found
    has_structure = False
    for indicator in section_indicators:
        if indicator.lower() in pdf_text.lower():
            has_structure = True
            if debug_mode:
                print(f"Found structure indicator: {indicator}")
            break

    if not has_structure:
        return False, "Document lacks organizational structure typically found in RFP documents."

    # If we got this far, the document seems to be an RFP
    return True, ""


def validate_rfp_inputs(data):
    """
    التحقق من جودة ومنطقية البيانات المدخلة لكتابة كراسة شروط فعالة.
    يستخدم تقنيات تحليل اللغة الطبيعية لتقييم البيانات.
    """
    # الحقول المطلوبة الأساسية
    required_field_names = [
        'competition_name',
        'competition_description',
        'competition_objectives',
        'government_entity'
    ]

    # التحقق من الحقول المفقودة
    missing_fields = [field for field in required_field_names if field not in data or not data[field]]
    if missing_fields:
        missing_str = "، ".join(missing_fields)
        return False, f"الحقول التالية مطلوبة ولكنها مفقودة أو فارغة: {missing_str}"

    # تحليل جودة البيانات المدخلة
    quality_issues = []

    # فحص اسم المنافسة
    comp_name = data.get('competition_name', '')
    if len(comp_name.split()) < 2:
        quality_issues.append("اسم المنافسة يبدو مختصراً جداً. الأسماء الفعالة عادة ما تكون وصفية ومحددة.")
    elif any(word.lower() in comp_name.lower() for word in ['test', 'تجربة', 'اختبار', 'xxx', 'asdf', 'hjkl']):
        quality_issues.append("اسم المنافسة يبدو أنه اختباري أو غير حقيقي.")

    # فحص وصف المنافسة
    comp_desc = data.get('competition_description', '')
    if len(comp_desc.split()) < 5:
        quality_issues.append("وصف المنافسة قصير جداً ولا يوفر معلومات كافية عن نطاق العمل.")
    elif comp_desc == comp_name or comp_desc == data.get('competition_objectives', ''):
        quality_issues.append(
            "وصف المنافسة يبدو مكرراً (مطابق لاسم المنافسة أو أهدافها). الوصف الفعال يقدم معلومات إضافية.")

    # فحص أهداف المنافسة
    comp_obj = data.get('competition_objectives', '')
    if len(comp_obj.split()) < 5:
        quality_issues.append("أهداف المنافسة قصيرة جداً. الأهداف الفعالة تكون واضحة ومحددة بشكل كافٍ.")
    elif not any(word in comp_obj for word in ['تحقيق', 'توفير', 'تطوير', 'تحسين', 'إنشاء', 'تنفيذ', 'زيادة', 'تعزيز']):
        quality_issues.append("أهداف المنافسة لا تبدو محددة بوضوح. الأهداف الفعالة عادة ما تتضمن أفعالاً محددة.")

    # فحص الجهة الحكومية
    gov_entity = data.get('government_entity', '')
    common_entities = ['وزارة', 'هيئة', 'مؤسسة', 'شركة', 'مديرية', 'إدارة', 'مجلس', 'مركز', 'جامعة', 'أمانة', 'بلدية']
    if not any(entity in gov_entity for entity in common_entities):
        quality_issues.append(
            "اسم الجهة الحكومية لا يتضمن كلمات معتادة مثل (وزارة، هيئة، مؤسسة...). تأكد من صحة اسم الجهة.")

    # فحص قيمة التكلفة إذا كانت موجودة
    if 'cost_value' in data and data['cost_value']:
        try:
            cost_str = str(data['cost_value']).strip()
            cost_str = re.sub(r'[^\d.]', '', cost_str)
            cost = float(cost_str) if cost_str else 0

            if cost == 0:
                quality_issues.append("قيمة التكلفة صفر، وهذا غير معتاد في كراسات الشروط الحقيقية.")
        except (ValueError, TypeError):
            quality_issues.append("قيمة التكلفة ليست رقماً صالحاً.")

    # فحص ترابط البيانات والتناسق
    if comp_name and comp_desc and comp_obj:
        # فحص التكرار الكامل بين الحقول
        if comp_name == comp_desc == comp_obj:
            quality_issues.append(
                "جميع الحقول الرئيسية متطابقة (الاسم، الوصف، الأهداف). هذا غير منطقي في البيانات الحقيقية.")

        # فحص وجود نص عشوائي أو غير ذي معنى
        random_texts = ['asdfgh', 'qwerty', '123456', 'test test', 'lorem ipsum', 'ابجد هوز', 'تجربة تجربة']
        if any(random_text in data.get(field, '').lower() for field in required_field_names for random_text in
               random_texts):
            quality_issues.append("تم اكتشاف نص عشوائي أو اختباري في البيانات المدخلة.")

    # التحقق من معلومات موقع التنفيذ إذا كانت موجودة
    location_fields = ['execution_city', 'execution_district', 'execution_region']
    if any(field in data and data[field] for field in location_fields):
        missing_location = [field for field in location_fields if field not in data or not data[field]]
        if missing_location:
            field_names = {
                'execution_city': 'المدينة',
                'execution_district': 'الحي',
                'execution_region': 'المنطقة'
            }
            missing_names = [field_names.get(field, field) for field in missing_location]
            missing_loc_str = "، ".join(missing_names)
            quality_issues.append(f"معلومات الموقع غير مكتملة. البيانات المفقودة: {missing_loc_str}")

    # إذا وجدت مشاكل في جودة البيانات، أعد تقريراً بذلك
    if quality_issues:
        combined_issues = "\n- ".join([""] + quality_issues)
        return False, f"البيانات المدخلة قد لا تكون كافية لإنشاء كراسة شروط فعالة لوجود المشاكل التالية:{combined_issues}"

    # إذا اجتازت جميع الفحوصات، فالبيانات تبدو منطقية وصالحة لإنشاء كراسة شروط فعالة
    return True, ""


def improved_rfp_with_validation(pdf_path, competition_name, competition_objectives, competition_description,
                                output_path, bypass_validation=True,vector_store = None,
 debug_mode=False, **kwargs):
    """
    Wrapper function to validate before improving an RFP document.
    """
    try:
        # Extract text from PDF
        print(f"📄 Processing user document: {pdf_path}")
        pdf_text = extract_text_from_pdf(pdf_path)

        # Try OCR if text extraction yielded insufficient text
        if not pdf_text or len(pdf_text.strip()) < 500:
            print("⚠️ Insufficient text extracted from PDF. Attempting OCR processing...")
            pdf_text = perform_ocr_on_pdf(pdf_path)

        # Final check on extracted text
        if not pdf_text or len(pdf_text.strip()) < 100:
            print("❌ System could not extract sufficient text from the file")
            return None

        # Verify document is an RFP (skip if bypass_validation is True)
        if not bypass_validation:
            # Debug mode will print detailed info about the validation process
            is_rfp, reason = is_valid_rfp_document(pdf_text, debug_mode=debug_mode)
            if not is_rfp:
                print(f"❌ The uploaded file doesn't appear to be an RFP document: {reason}")
                return None
        else:
            print("ℹ️ Document type validation bypassed")

        # Validate user input
        user_data = {
            'competition_name': competition_name,
            'competition_objectives': competition_objectives,
            'competition_description': competition_description,
            'government_entity': kwargs.get('government_entity', ''),
        }
        user_data.update(kwargs)

        inputs_valid, validation_reason = validate_rfp_inputs(user_data)
        if not inputs_valid:
            print(f"❌ AI Analysis: Input data is insufficient for creating an effective RFP: {validation_reason}")
            return None

        # If all validations passed, proceed with improving the RFP
        print("✅ Document and data validated. Starting improvement process...")

        # Call the original improvement function
        return improve_rfp_with_extracted_text(
            pdf_path, competition_name, competition_objectives, competition_description,
            output_path, vector_store, **kwargs
        )

    except Exception as e:
        print(f"❌ Error during RFP validation: {str(e)}")
        import traceback
        traceback.print_exc()
        return None
def generate_rfp_with_validation(
        competition_name, competition_objectives, competition_description, government_entity,
        cost_value, cost_method, start_stage, end_stage, technical_docs, alternative_offers,
        initial_guarantee, pause_period, penalties, execution_city, execution_district, execution_region,
        required_materials, scope_summary, special_terms, output_dir, static_dir
):
    """
    Wrapper function to validate before generating a new RFP document.
    """
    try:
        # Collect user input data for validation
        user_data = {
            'competition_name': competition_name,
            'competition_objectives': competition_objectives,
            'competition_description': competition_description,
            'government_entity': government_entity,
            'cost_value': cost_value,
            'cost_method': cost_method,
            'start_stage': start_stage,
            'end_stage': end_stage,
            'technical_docs': technical_docs,
            'alternative_offers': alternative_offers,
            'initial_guarantee': initial_guarantee,
            'pause_period': pause_period,
            'penalties': penalties,
            'execution_city': execution_city,
            'execution_district': execution_district,
            'execution_region': execution_region,
            'required_materials': required_materials,
            'scope_summary': scope_summary,
            'special_terms': special_terms
        }

        # Validate user input
        inputs_valid, validation_reason = validate_rfp_inputs(user_data)
        if not inputs_valid:
            print(f"❌ AI Analysis: Input data is insufficient for creating an effective RFP: {validation_reason}")
            return None

        # Supplementary checks for RFP generation
        if not scope_summary or len(scope_summary.strip()) < 20:
            # Generate default summary from objectives and description
            scope_summary = f"The scope of work includes implementing {competition_description} with the goal of achieving {competition_objectives}."
            print(f"⚠️ Generated default scope summary: {scope_summary}")

        if not special_terms or len(special_terms.strip()) < 10:
            special_terms = "No additional special terms."
            print(f"⚠️ Set default value for special terms: {special_terms}")

        # If all validations passed, proceed with generating the RFP
        print("✅ Data validated. Starting RFP generation...")

        # Call the original generation function
        return generate_rfp_document(
            competition_name, competition_objectives, competition_description,
            government_entity, cost_value, cost_method, start_stage, end_stage,
            technical_docs, alternative_offers, initial_guarantee, pause_period,
            penalties, execution_city, execution_district, execution_region,
            required_materials, scope_summary, special_terms, output_dir, static_dir
        )

    except Exception as e:
        print(f"❌ Error during RFP data validation: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


